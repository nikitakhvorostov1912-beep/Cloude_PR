"""
Регенерация стенограммы встречи 2 (17.03.2026) с полным текстом (501 строка).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def set_run_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def create_summary_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, value) in enumerate(rows_data):
        cell_key = table.cell(i, 0)
        cell_val = table.cell(i, 1)
        cell_key.width = Cm(5)
        cell_val.width = Cm(12)
        set_cell_shading(cell_key, '2B579A')
        p = cell_key.paragraphs[0]
        run = p.add_run(key)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p2 = cell_val.paragraphs[0]
        run2 = p2.add_run(value)
        set_run_font(run2, size=10)
    return table


def create_stats_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(['Параметр', 'Значение']):
        cell = table.cell(0, j)
        set_cell_shading(cell, '2B579A')
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
    for i, (key, value) in enumerate(rows_data):
        table.cell(i + 1, 0).paragraphs[0].add_run(key).font.size = Pt(10)
        table.cell(i + 1, 1).paragraphs[0].add_run(str(value)).font.size = Pt(10)
    return table


def read_transcript(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return [line.strip() for line in f.readlines() if line.strip()]


def split_into_paragraphs(lines, chunk_size=6):
    paragraphs = []
    for i in range(0, len(lines), chunk_size):
        chunk = lines[i:i + chunk_size]
        paragraphs.append(' '.join(chunk))
    return paragraphs


def count_words(lines):
    return len(' '.join(lines).split())


# ============================================================

m2_lines = read_transcript(r'D:\Cloude_PR\docs\meetings\temp\m2_text_lines.txt')
print(f'Прочитано строк: {len(m2_lines)}')

m2_chapters = [
    (
        'Открытие встречи и обзор результатов предыдущей встречи',
        [
            'Обзор шести выделенных задач по результатам предыдущей встречи',
            'Проблемы со связью у Дениса (отключение интернета и света)',
            'Демонстрация документа с задачами в режиме редактирования',
        ],
        (0, 18),
    ),
    (
        'Доработка карточки номенклатуры и настройки поставщиков',
        [
            'Уже существуют свойства номенклатуры: размер, плотность, объем',
            'Необходимо добавить: вес, срок производства, срок доставки',
            'Связка номенклатура-поставщик — основа для планирования',
            'Загрузка параметров из Excel с проверкой или ручное заполнение через форму',
            'Кратность и минимальная партия — уточнить единицы (упаковки vs штуки)',
            'Поля универсальные для российских и ВЭД-поставщиков',
        ],
        (18, 68),
    ),
    (
        'Рабочее место планирования закупок ВЭД: отдельная вкладка или инструмент',
        [
            'Копируем текущее рабочее место, выносим ВЭД на отдельную вкладку',
            'Валентин предлагает отдельную гиперссылку вместо вкладки',
            'Отборы: по номенклатурам с признаком ГТД, по заводам, по складам',
            'Аналоги номенклатур и остатки по определенным складам',
            'Права доступа для разделения российских и ВЭД-закупщиков',
            'Детализация: возможность провалиться из поля «Едет» в накладную и контейнер',
        ],
        (68, 130),
    ),
    (
        'Цветовая индикация и детализация статусов',
        [
            'Текущая проблема: при наличии хотя бы одного неподтвержденного заказа вся сумма синяя',
            'Плюсиками (50 + 20) реализовать не удалось — программист не смог',
            'Решено: не перегружать форму цветами, оставить текущую индикацию',
            'Использовать фильтры по статусам вместо визуальной перегрузки',
            'Статусы ВЭД: в пути, в порту, подготавливается к отправке и т.д.',
        ],
        (130, 190),
    ),
    (
        'Архитектура документооборота: план закупок, заказы, накладные ВЭД',
        [
            'Цепочка: план производства -> потребности -> план закупок -> заказы -> накладная ВЭД',
            'Накладная ВЭД — не альтернатива плану закупок, а последующий этап',
            'Оставляем документ «План закупок» как основу, расширяем его функциональность',
            'Обсуждение: один план закупок с множеством строк vs много документов по одной строке',
            'Согласование с производством в Китае — отдельный этап между планом и заказом',
            'Группировка планов закупок по поставщику со всеми номенклатурами',
        ],
        (190, 270),
    ),
    (
        'Планы закупок по поставщикам, статусы и производительность',
        [
            'Формирование планов закупок по поставщику с указанием всех номенклатур',
            'Статусы: привязаны к документу «План закупок» — синий (подготовка) и зеленый (утвержден)',
            'Проблема производительности MISA при работе с большими документами',
            'Не тащить планы закупок из MISA в ERP — оставить в MISA, в ERP только заказы',
            'Массовое формирование планов закупок одной кнопкой',
        ],
        (270, 330),
    ),
    (
        'Плановая себестоимость: три уровня расчета',
        [
            'Первый уровень: полное наполнение контейнера одной позицией (20-фут и 40-фут)',
            'Нюанс: цена ФОБ не равна конечной себестоимости из-за различий упаковки и габаритов',
            'Сравнение с российскими ценами для оценки рентабельности',
            'Подсветка оптимального поставщика по цене и срокам',
            'Возможность редактировать количество и менять поставщиков',
        ],
        (330, 377),
    ),
    (
        'Групповое планирование и динамическая себестоимость',
        [
            'Первая себестоимость — индивидуальная по позиции, не используется при групповом планировании',
            'Вторая себестоимость — при формировании заказа поставщику (большой список по нажатию кнопки)',
            'Корректировка количества вручную: дополнение контейнера до полного наполнения',
            'Динамический пересчет себестоимости при изменении количества позиций',
            'Согласование: Виктория проверяет наполнение и себестоимость, утверждает -> заказы в ERP',
        ],
        (377, 402),
    ),
    (
        'Актуализация данных для российских закупщиков',
        [
            'Проблема: российские закупщики видят устаревшие сроки поставки из Китая',
            'Валентин формирует заказы на 4 месяца вперед, закупщики смотрят максимум на месяц',
            'Решение: закупщики видят только реально отгруженные накладные с актуальными датами',
            'Частичные поставки: из 30 000 единиц может прийти сначала 10 000, остальные позже',
            'Приобретения из накладных сокращают «Едет» и пополняют остатки автоматически',
            'Итого: пока не трогать рабочее место для России, вернуться после реализации ВЭД-блока',
        ],
        (402, 455),
    ),
    (
        'Фрахт, логистика и расчет себестоимости',
        [
            'Фрахт реальный (при отгрузке) vs плановый (от перевозчиков перед отгрузкой)',
            'Таблица фрахта заполняется помесячно, берется последнее актуальное значение',
            'Логистика — существенная часть себестоимости, программа должна корректно учитывать долю в пошлине',
            'Проблема страховки: курсы валют ежедневно разные, небольшие расхождения неизбежны',
        ],
        (455, 478),
    ),
    (
        'Итоги и следующие шаги',
        [
            'Денис до конца недели актуализирует файл с задачами и дополнит обсужденным',
            'Участники должны подумать: какие данные нужны для расчета первой себестоимости',
            'Помимо веса и объема — возможно добавить дополнительные параметры по умолчанию для расчета',
            'Обратная связь в понедельник, затем запланировать финальную встречу',
            'Цель: финализировать ТЗ и перейти к разработке',
        ],
        (478, 501),
    ),
]

m2_decisions = [
    'Доработать карточку номенклатуры: использовать существующие свойства + добавить вес, сроки',
    'Параметры привязать к связке номенклатура-поставщик (не к самой номенклатуре)',
    'Сделать универсальные поля для российских и ВЭД-поставщиков (срок производства, доставки)',
    'Рабочее место ВЭД — отдельная гиперссылка (не вкладка в общем рабочем месте)',
    'Цветовую индикацию оставить текущую, не добавлять третий цвет для импорта',
    'Использовать фильтры по статусам в отборах вместо визуальной перегрузки',
    'Оставить документ «План закупок» как основу, расширить его для ВЭД',
    'Массовое формирование планов закупок одной кнопкой с учетом всех параметров',
    'Три уровня себестоимости: при планировании, при заказе, в накладной',
    'Динамический пересчет себестоимости при корректировке количества в заказе',
    'Фрахт учитывать через помесячную таблицу, брать последнее актуальное значение',
    'Рабочее место для России пока не трогать — вернуться после реализации ВЭД',
    'Денис до конца недели актуализирует файл с задачами, обратная связь в понедельник',
    'Уточнить единицы для кратности и минимальной партии (упаковки или штуки)',
]

# ============================================================
# Генерация документа
# ============================================================

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in ['Heading 1', 'Heading 2']:
    s = doc.styles[level]
    s.font.name = 'Calibri'
    s.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

# --- Титульная страница ---
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('СТЕНОГРАММА')
set_run_font(run, size=28, bold=True, color=(0x2B, 0x57, 0x9A))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('ВИДЕОВСТРЕЧИ 17 МАРТА 2026 Г.')
set_run_font(run, size=20, bold=True, color=(0x2B, 0x57, 0x9A))

doc.add_paragraph()

topic_p = doc.add_paragraph()
topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = topic_p.add_run('Продолжение проработки рабочего места ВЭД — детализация задач')
set_run_font(run, size=14, color=(0x55, 0x55, 0x55))

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Длительность: 82 минуты')
set_run_font(run, size=12, color=(0x77, 0x77, 0x77))

doc.add_page_break()

# --- Резюме ---
doc.add_heading('Общая информация', level=1)

participants = [
    'Денис (аналитик, ведущий)',
    'Валентин (менеджер ВЭД)',
    'Иван (разработчик, Интерсофт)',
    'Виктория Валерьевна (руководитель закупок, упоминается)',
]

main_topics = [
    'Продолжение проработки рабочего места ВЭД — детализация задач',
    'Доработка карточки номенклатуры и настройки поставщиков',
    'Архитектура документооборота: план закупок, заказы, накладные',
    'Цветовая индикация и фильтрация статусов',
    'Плановая себестоимость и наполнение контейнеров',
    'Фрахт, логистика и страховка в расчете себестоимости',
    'Актуализация данных для российских закупщиков',
]

summary_data = [
    ('Дата встречи', '17 марта 2026 г.'),
    ('Длительность', '82 минуты'),
    ('Участники', ', '.join(participants)),
    ('Основные темы', '\n'.join(f'- {t}' for t in main_topics)),
]
create_summary_table(doc, summary_data)

doc.add_paragraph()

# --- Ключевые решения ---
doc.add_heading('Ключевые решения и договоренности', level=1)
for i, d in enumerate(m2_decisions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    set_run_font(run, bold=True)
    run2 = p.add_run(d)
    set_run_font(run2, size=11)

doc.add_page_break()

# --- Главы ---
doc.add_heading('Содержание встречи', level=1)
doc.add_paragraph()

# Оглавление
doc.add_heading('Оглавление', level=2)
for i, (ch_title, _, _) in enumerate(m2_chapters, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {ch_title}')
    set_run_font(run, size=11, color=(0x2B, 0x57, 0x9A))
doc.add_paragraph()

# Главы с содержимым
for i, (ch_title, ch_key_points, ch_line_range) in enumerate(m2_chapters, 1):
    doc.add_heading(f'{i}. {ch_title}', level=1)

    doc.add_heading('Ключевые тезисы', level=2)
    for kp in ch_key_points:
        p = doc.add_paragraph()
        run = p.add_run(f'\u2022 {kp}')
        set_run_font(run, size=10, color=(0x44, 0x44, 0x44))

    doc.add_heading('Текст стенограммы', level=2)
    start, end = ch_line_range
    chunk_lines = m2_lines[start:end]
    paragraphs = split_into_paragraphs(chunk_lines, chunk_size=6)
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_run_font(run, size=10, color=(0x33, 0x33, 0x33))

doc.add_page_break()

# --- Статистика ---
doc.add_heading('Статистика встречи', level=1)

word_count = count_words(m2_lines)
stats_data = [
    ('Общая длительность', '82 минуты'),
    ('Количество строк транскрипта', str(len(m2_lines))),
    ('Количество слов', f'{word_count:,}'.replace(',', ' ')),
    ('Количество глав', str(len(m2_chapters))),
    ('Количество ключевых решений', str(len(m2_decisions))),
    ('Участников', str(len(participants))),
]
create_stats_table(doc, stats_data)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- Конец стенограммы ---')
set_run_font(run, size=10, color=(0x99, 0x99, 0x99))

# Сохранение
output_path = r'D:\Cloude_PR\docs\meetings\стенограмма-2026-03-17.docx'
doc.save(output_path)
print(f'Создан: {output_path}')
print(f'Строк: {len(m2_lines)}, Слов: {word_count}, Глав: {len(m2_chapters)}, Решений: {len(m2_decisions)}')
