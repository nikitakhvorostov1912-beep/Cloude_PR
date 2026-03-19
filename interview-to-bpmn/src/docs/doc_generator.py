"""Generate Word documents from process data using templates."""
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


def generate_documents(processes: dict, project_dir: str, config: dict) -> list:
    """Generate Word documents for all processes.

    Args:
        processes: Extracted processes dict.
        project_dir: Project directory path.
        config: Application config.

    Returns:
        List of generated document file paths.
    """
    output_dir = Path(project_dir) / "output"
    os.makedirs(output_dir, exist_ok=True)

    doc_files = []

    # 1. Generate survey report
    report_path = _generate_survey_report(processes, output_dir, config)
    doc_files.append(report_path)

    # 2. Generate process cards
    for proc in processes.get("processes", []):
        card_path = _generate_process_card(proc, output_dir, config)
        doc_files.append(card_path)

    return doc_files


def _generate_survey_report(processes: dict, output_dir: Path, config: dict) -> str:
    """Generate survey/investigation report per GOST 34.602-89 structure."""
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Title page
    _add_title_page(doc, processes)

    # Table of contents placeholder
    doc.add_page_break()
    doc.add_heading("СОДЕРЖАНИЕ", level=1)
    doc.add_paragraph("(Автоматически обновить после открытия в Word)")

    # 1. General information
    doc.add_page_break()
    doc.add_heading("1. ОБЩИЕ СВЕДЕНИЯ", level=1)

    dept = processes.get("department", "Не указан")
    doc.add_heading("1.1. Объект обследования", level=2)
    doc.add_paragraph(f"Отдел/подразделение: {dept}")
    doc.add_paragraph(f"Респондент: {processes.get('respondent', 'Не указан')}")

    meta = processes.get("transcript_metadata", {})
    if meta:
        doc.add_paragraph(f"Длительность интервью: {meta.get('total_duration_formatted', 'Н/Д')}")
        doc.add_paragraph(f"Количество спикеров: {meta.get('speaker_count', 'Н/Д')}")

    doc.add_heading("1.2. Цель обследования", level=2)
    doc.add_paragraph(
        "Целью предпроектного обследования является анализ текущих бизнес-процессов "
        "подразделения (AS IS), выявление проблемных зон и формирование требований "
        "к целевой автоматизированной системе (TO BE)."
    )

    # 2. Business processes
    doc.add_page_break()
    doc.add_heading("2. ОПИСАНИЕ БИЗНЕС-ПРОЦЕССОВ", level=1)

    for i, proc in enumerate(processes.get("processes", []), 1):
        doc.add_heading(f"2.{i}. {proc.get('name', 'Процесс')}", level=2)

        # Process overview table
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        rows_data = [
            ("Триггер", proc.get("trigger", "Не определён")),
            ("Результат", proc.get("result", "Не определён")),
            ("Периодичность", proc.get("frequency", "Не определена")),
            ("Тип", "AS IS" if proc.get("type") == "as_is" else "TO BE"),
            ("Участники", ", ".join(
                p.get("role", "") for p in proc.get("participants", [])
            ) or "Не определены"),
            ("Системы", ", ".join(
                i.get("system", "") for i in proc.get("integrations", [])
            ) or "Не определены"),
        ]

        for j, (label, value) in enumerate(rows_data):
            table.rows[j].cells[0].text = label
            table.rows[j].cells[1].text = value
            # Bold label
            for paragraph in table.rows[j].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Steps
        steps = proc.get("steps", [])
        if steps:
            doc.add_heading("Шаги процесса:", level=3)
            for k, step in enumerate(steps, 1):
                performer = step.get("performer", "")
                performer_str = f" [{performer}]" if performer else ""
                doc.add_paragraph(
                    f"{k}. {step.get('name', 'Шаг')}{performer_str}",
                    style="List Number",
                )
                if step.get("description"):
                    doc.add_paragraph(f"   {step['description']}")

        # Decisions
        decisions = proc.get("decisions", [])
        if decisions:
            doc.add_heading("Точки принятия решений:", level=3)
            for dec in decisions:
                doc.add_paragraph(f"Вопрос: {dec.get('question', '?')}")
                for opt in dec.get("options", []):
                    doc.add_paragraph(
                        f"  - {opt.get('condition', '')}: → {opt.get('next_step', '')}",
                        style="List Bullet",
                    )

        # Documents
        all_docs_in = set()
        all_docs_out = set()
        for step in steps:
            all_docs_in.update(step.get("documents_in", []))
            all_docs_out.update(step.get("documents_out", []))

        if all_docs_in or all_docs_out:
            doc.add_heading("Документооборот:", level=3)
            if all_docs_in:
                doc.add_paragraph(f"Входящие: {', '.join(all_docs_in)}")
            if all_docs_out:
                doc.add_paragraph(f"Исходящие: {', '.join(all_docs_out)}")

    # 3. Pain points
    all_pain_points = []
    for proc in processes.get("processes", []):
        for pp in proc.get("pain_points", []):
            all_pain_points.append({**pp, "process": proc.get("name", "")})

    if all_pain_points or processes.get("general_issues"):
        doc.add_page_break()
        doc.add_heading("3. ВЫЯВЛЕННЫЕ ПРОБЛЕМЫ", level=1)

        if all_pain_points:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Процесс"
            hdr[1].text = "Проблема"
            hdr[2].text = "Влияние"
            hdr[3].text = "Критичность"
            for cell in hdr:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

            for pp in all_pain_points:
                row = table.add_row().cells
                row[0].text = pp.get("process", "")
                row[1].text = pp.get("description", "")
                row[2].text = pp.get("impact", "")
                row[3].text = pp.get("severity", "")

        for issue in processes.get("general_issues", []):
            doc.add_paragraph(
                f"[{issue.get('category', 'Общее')}] {issue.get('description', '')}",
                style="List Bullet",
            )

    # 4. Automation requests
    auto_requests = processes.get("automation_requests", [])
    if auto_requests:
        doc.add_page_break()
        doc.add_heading("4. ТРЕБОВАНИЯ К АВТОМАТИЗАЦИИ", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "№"
        hdr[1].text = "Требование"
        hdr[2].text = "Приоритет"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for i, req in enumerate(auto_requests, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = req.get("description", "")
            row[2].text = req.get("priority", "")

    # Save
    filename = f"survey_report_{processes.get('department', 'dept')}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))

    return str(filepath)


def _add_title_page(doc: Document, processes: dict):
    """Add GOST-style title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ ПО ПРЕДПРОЕКТНОМУ ОБСЛЕДОВАНИЮ")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept = processes.get("department", "подразделения")
    run = subtitle.add_run(f"Отдел: {dept}")
    run.font.size = Pt(14)

    doc.add_paragraph("")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%d.%m.%Y"))
    run.font.size = Pt(12)


def _generate_process_card(process: dict, output_dir: Path, config: dict) -> str:
    """Generate a process card document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    proc_name = process.get("name", "Процесс")
    proc_id = process.get("id", "proc")

    # Header
    doc.add_heading("КАРТОЧКА БИЗНЕС-ПРОЦЕССА", level=1)
    doc.add_heading(proc_name, level=2)

    # Main info table
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    info_rows = [
        ("Идентификатор", proc_id),
        ("Тип", "AS IS" if process.get("type") == "as_is" else "TO BE"),
        ("Триггер", process.get("trigger", "")),
        ("Результат", process.get("result", "")),
        ("Периодичность", process.get("frequency", "")),
        ("Участники", ", ".join(p.get("role", "") for p in process.get("participants", []))),
        ("Документы (вход)", ", ".join(
            d for s in process.get("steps", []) for d in s.get("documents_in", [])
        )),
        ("Документы (выход)", ", ".join(
            d for s in process.get("steps", []) for d in s.get("documents_out", [])
        )),
    ]

    for i, (label, value) in enumerate(info_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    # Steps
    doc.add_heading("Шаги процесса", level=2)
    for i, step in enumerate(process.get("steps", []), 1):
        doc.add_paragraph(
            f"{i}. {step.get('name', '')} [{step.get('performer', '')}]",
            style="List Number",
        )

    # Pain points
    pain_points = process.get("pain_points", [])
    if pain_points:
        doc.add_heading("Проблемные зоны", level=2)
        for pp in pain_points:
            doc.add_paragraph(
                f"[{pp.get('severity', 'medium')}] {pp.get('description', '')}",
                style="List Bullet",
            )

    # BPMN diagram placeholder
    doc.add_heading("Схема процесса (BPMN 2.0)", level=2)

    # Try to embed BPMN image
    bpmn_image = output_dir / f"{proc_id}_overview.png"
    if bpmn_image.exists():
        doc.add_picture(str(bpmn_image), width=Inches(6))
    else:
        doc.add_paragraph("[Схема будет добавлена после генерации BPMN]")

    # Save
    filename = f"process_card_{proc_id}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))

    return str(filepath)
