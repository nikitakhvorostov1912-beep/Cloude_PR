"""Скрипт генерации ВСЕХ документов для тестового проекта.

Выполняет полный цикл:
1. Загружает процессы из _all_processes.json
2. Конвертирует process → BPMN JSON
3. Генерирует BPMN XML + SVG
4. Генерирует Visio (.vsdx) для каждого процесса
5. Генерирует Word (описание процессов + лист требований)
6. Генерирует Excel (требования + GAP-анализ)

Запуск:
    cd backend
    .venv/Scripts/python scripts/generate_all_docs.py
"""

import json
import logging
import sys
from pathlib import Path
from typing import Any

# Добавляем backend в путь
backend_dir = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(backend_dir))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("generate_docs")

# Ищем тестовый проект
DATA_DIR = backend_dir / "data" / "projects"


def find_test_project() -> Path | None:
    """Находит проект с _all_processes.json."""
    if not DATA_DIR.is_dir():
        return None
    for d in sorted(DATA_DIR.iterdir()):
        if d.is_dir() and (d / "processes" / "_all_processes.json").is_file():
            return d
    # Fallback: любой проект с директорией processes
    for d in sorted(DATA_DIR.iterdir()):
        if d.is_dir() and (d / "processes").is_dir():
            return d
    return None


def load_json(path: Path) -> list | dict | None:
    """Загружает JSON-файл."""
    if not path.is_file():
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(path: Path, data: Any) -> None:
    """Сохраняет JSON-файл."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def main() -> None:
    project_dir = find_test_project()
    if not project_dir:
        logger.error("Тестовый проект не найден в %s", DATA_DIR)
        sys.exit(1)

    project_id = project_dir.name
    logger.info("Проект: %s", project_id)

    processes_dir = project_dir / "processes"
    bpmn_dir = project_dir / "bpmn"
    output_dir = project_dir / "output"
    visio_dir = project_dir / "visio"
    bpmn_dir.mkdir(exist_ok=True)
    output_dir.mkdir(exist_ok=True)
    visio_dir.mkdir(exist_ok=True)

    # =================================================================
    # 1. Загрузка процессов
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 1: Загрузка процессов")
    logger.info("=" * 60)

    processes = load_json(processes_dir / "_all_processes.json")
    if not processes or not isinstance(processes, list):
        logger.error("Файл _all_processes.json не найден или пуст")
        sys.exit(1)

    logger.info("Загружено процессов: %d", len(processes))

    # =================================================================
    # 2. Конвертация process → BPMN JSON
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 2: Конвертация process → BPMN JSON")
    logger.info("=" * 60)

    from app.bpmn.process_to_bpmn import ProcessToBpmnConverter

    converter = ProcessToBpmnConverter()
    bpmn_jsons: list[dict] = []

    for proc in processes:
        bpmn_json = converter.convert(proc)
        bpmn_jsons.append(bpmn_json)

        # Сохраняем BPMN JSON
        pid = bpmn_json["process_id"]
        save_json(processes_dir / f"{pid}_bpmn.json", bpmn_json)
        logger.info(
            "  [✓] %s: %d элементов, %d потоков",
            pid, len(bpmn_json["elements"]), len(bpmn_json["flows"]),
        )

    # =================================================================
    # 3. Генерация BPMN XML + SVG
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 3: Генерация BPMN 2.0 XML + SVG")
    logger.info("=" * 60)

    from app.bpmn.json_to_bpmn import BpmnConverter
    from app.bpmn.layout import BpmnLayout

    bpmn_converter = BpmnConverter()
    layout_engine = BpmnLayout()

    for bpmn_json in bpmn_jsons:
        pid = bpmn_json["process_id"]
        pname = bpmn_json["process_name"]

        try:
            # Layout
            layout = layout_engine.calculate_layout(bpmn_json)
            bpmn_json["layout"] = layout

            # BPMN XML
            bpmn_xml = bpmn_converter.convert(bpmn_json)
            bpmn_path = bpmn_dir / f"{pid}.bpmn"
            bpmn_path.write_text(bpmn_xml, encoding="utf-8")
            logger.info("  [✓] BPMN XML: %s (%.1f KB)", bpmn_path.name, len(bpmn_xml) / 1024)

            # SVG
            from app.services.bpmn_service import BPMNService
            svg_content = BPMNService._render_svg(bpmn_json, layout)
            svg_path = bpmn_dir / f"{pid}.svg"
            svg_path.write_text(svg_content, encoding="utf-8")
            logger.info("  [✓] SVG: %s (%.1f KB)", svg_path.name, len(svg_content) / 1024)

        except Exception as exc:
            logger.error("  [✗] Ошибка BPMN для %s: %s", pid, exc)

    # =================================================================
    # 4. Генерация Visio (.vsdx)
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 4: Генерация Visio (.vsdx)")
    logger.info("=" * 60)

    try:
        from app.visio import generate_visio

        for bpmn_json in bpmn_jsons:
            pid = bpmn_json["process_id"]
            pname = bpmn_json["process_name"]
            visio_path = visio_dir / f"{pid}.vsdx"

            try:
                generate_visio(bpmn_json, visio_path)
                size_kb = visio_path.stat().st_size / 1024
                logger.info("  [✓] Visio: %s (%.1f KB) — %s", visio_path.name, size_kb, pname)
            except Exception as exc:
                logger.error("  [✗] Ошибка Visio для %s: %s", pid, exc)

    except ImportError as exc:
        logger.warning("  [!] Visio-генератор недоступен: %s", exc)

    # =================================================================
    # 5. Генерация Word-документов
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 5: Генерация Word-документов")
    logger.info("=" * 60)

    # Загружаем мета-данные проекта
    project_meta = load_json(project_dir / "project.json") or {}
    project_name = project_meta.get("name", "Тестовый проект")

    # 5a. Описание процессов
    try:
        from app.docs.process_doc import ProcessDocGenerator

        doc_gen = ProcessDocGenerator(company=project_name)
        doc_path = output_dir / "описание_процессов.docx"
        doc_gen.generate(processes=processes, output_path=doc_path, project_name=project_name)
        size_kb = doc_path.stat().st_size / 1024
        logger.info("  [✓] Описание процессов: %s (%.1f KB)", doc_path.name, size_kb)
    except Exception as exc:
        logger.error("  [✗] Ошибка Word (процессы): %s", exc)

    # 5b. Лист требований (Word)
    requirements = load_json(processes_dir / "_requirements.json")
    if requirements:
        try:
            from app.docs.doc_generator import DocGenerator, format_date_russian, safe_str
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt

            gen = ProcessDocGenerator(company=project_name)
            doc = gen._create_document(title=f"Лист требований — {project_name}")
            gen._add_title_page(doc, "Лист требований", project_name)
            gen._add_table_of_contents(doc)
            gen._add_header(doc, f"{project_name} — Лист требований")
            gen._add_footer_with_pages(doc)

            req_items = requirements if isinstance(requirements, list) else requirements.get("requirements", [])

            by_type: dict[str, list] = {}
            for req in req_items:
                rtype = req.get("type", "Прочие")
                by_type.setdefault(rtype, []).append(req)

            type_names = {
                "FR": "Функциональные требования",
                "NFR": "Нефункциональные требования",
                "IR": "Интеграционные требования",
            }

            for rtype, reqs in by_type.items():
                section_name = type_names.get(rtype, rtype)
                doc.add_heading(section_name, level=1)

                for req in reqs:
                    req_id = req.get("id", "")
                    req_name = req.get("name", "")
                    doc.add_heading(f"{req_id}: {req_name}", level=2)

                    desc = req.get("description", "")
                    if desc:
                        doc.add_paragraph(desc)

                    fields = [
                        ("Приоритет", req.get("priority", "")),
                        ("Категория", req.get("category", "")),
                        ("Источник", req.get("source", "")),
                        ("Трудоёмкость", f"{req.get('effort_hours', '—')} ч"),
                        ("Подсистема 1С", req.get("erp_subsystem", "")),
                    ]

                    table = doc.add_table(rows=len(fields), cols=2)
                    table.style = "Table Grid"
                    for row, (label, value) in zip(table.rows, fields):
                        row.cells[0].text = label
                        row.cells[1].text = safe_str(value)
                    gen._format_table_header(table)

                    criteria = req.get("acceptance_criteria", [])
                    if criteria:
                        doc.add_heading("Критерии приёмки", level=3)
                        for criterion in criteria:
                            doc.add_paragraph(str(criterion), style="List Bullet")

            req_doc_path = output_dir / "лист_требований.docx"
            doc.save(str(req_doc_path))
            size_kb = req_doc_path.stat().st_size / 1024
            logger.info("  [✓] Лист требований: %s (%.1f KB)", req_doc_path.name, size_kb)
        except Exception as exc:
            logger.error("  [✗] Ошибка Word (требования): %s", exc)

    # =================================================================
    # 6. Генерация Excel-документов
    # =================================================================
    logger.info("=" * 60)
    logger.info("Шаг 6: Генерация Excel-документов")
    logger.info("=" * 60)

    # 6a. Требования Excel
    if requirements:
        try:
            from app.docs.doc_generator import (
                EXCEL_ALIGNMENT_LEFT,
                EXCEL_BORDER,
                EXCEL_FILLS,
                EXCEL_FONTS,
                DocGenerator,
            )
            from openpyxl import Workbook

            req_items = requirements if isinstance(requirements, list) else requirements.get("requirements", [])

            wb = Workbook()
            ws = wb.active
            ws.title = "Требования"

            headers = [
                "ID", "Тип", "Название", "Описание", "Категория",
                "Приоритет", "Источник", "Критерии приёмки",
                "Трудоёмкость (ч)", "Подсистема 1С",
            ]
            DocGenerator._write_excel_header(ws, headers)

            for row_idx, req in enumerate(req_items, start=2):
                ws.cell(row=row_idx, column=1, value=req.get("id", ""))
                ws.cell(row=row_idx, column=2, value=req.get("type", ""))
                ws.cell(row=row_idx, column=3, value=req.get("name", ""))
                ws.cell(row=row_idx, column=4, value=req.get("description", ""))
                ws.cell(row=row_idx, column=5, value=req.get("category", ""))
                ws.cell(row=row_idx, column=6, value=req.get("priority", ""))
                ws.cell(row=row_idx, column=7, value=req.get("source", ""))

                criteria = req.get("acceptance_criteria", [])
                if isinstance(criteria, list):
                    criteria = "\n".join(f"- {c}" for c in criteria)
                ws.cell(row=row_idx, column=8, value=str(criteria))

                ws.cell(row=row_idx, column=9, value=req.get("effort_hours", ""))
                ws.cell(row=row_idx, column=10, value=req.get("erp_subsystem", ""))

                priority = str(req.get("priority", "")).lower()
                if priority in EXCEL_FILLS:
                    ws.cell(row=row_idx, column=6).fill = EXCEL_FILLS[priority]

                for col_idx in range(1, len(headers) + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.border = EXCEL_BORDER
                    cell.alignment = EXCEL_ALIGNMENT_LEFT
                    cell.font = EXCEL_FONTS["normal"]

            DocGenerator._auto_adjust_columns(ws)
            DocGenerator._freeze_panes(ws)
            DocGenerator._add_autofilter(ws)

            req_excel_path = output_dir / "требования.xlsx"
            wb.save(str(req_excel_path))
            size_kb = req_excel_path.stat().st_size / 1024
            logger.info("  [✓] Требования Excel: %s (%.1f KB)", req_excel_path.name, size_kb)
        except Exception as exc:
            logger.error("  [✗] Ошибка Excel (требования): %s", exc)

    # 6b. GAP-анализ Excel
    gaps = load_json(processes_dir / "_gap_analysis.json")
    if gaps and isinstance(gaps, list):
        try:
            from app.docs.doc_generator import (
                EXCEL_ALIGNMENT_LEFT,
                EXCEL_BORDER,
                EXCEL_FILLS,
                EXCEL_FONTS,
                DocGenerator,
            )
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "GAP-анализ"

            headers = [
                "Процесс", "Шаг №", "Шаг", "Покрытие",
                "Механизм 1С", "Разрыв (GAP)", "Рекомендация",
                "Трудоёмкость", "Приоритет", "Риски",
            ]
            DocGenerator._write_excel_header(ws, headers)

            row_idx = 2
            for gap in gaps:
                process_name = gap.get("process_name", gap.get("process_id", ""))
                step_analysis = gap.get("step_analysis", [])

                for step in step_analysis:
                    ws.cell(row=row_idx, column=1, value=process_name)
                    ws.cell(row=row_idx, column=2, value=step.get("step_order", ""))
                    ws.cell(row=row_idx, column=3, value=step.get("step_name", ""))
                    ws.cell(row=row_idx, column=4, value=step.get("coverage", ""))
                    ws.cell(row=row_idx, column=5, value=step.get("erp_mechanism", ""))
                    ws.cell(row=row_idx, column=6, value=step.get("gap_description", ""))
                    ws.cell(row=row_idx, column=7, value=step.get("recommendation", ""))
                    ws.cell(row=row_idx, column=8, value=step.get("effort", ""))
                    ws.cell(row=row_idx, column=9, value=step.get("priority", ""))

                    risks = step.get("risks", [])
                    if isinstance(risks, list):
                        risks = "; ".join(risks)
                    ws.cell(row=row_idx, column=10, value=str(risks))

                    coverage = str(step.get("coverage", "")).lower()
                    coverage_colors = {
                        "full": "green", "partial": "yellow",
                        "custom": "should", "absent": "red",
                    }
                    fill_key = coverage_colors.get(coverage)
                    if fill_key and fill_key in EXCEL_FILLS:
                        ws.cell(row=row_idx, column=4).fill = EXCEL_FILLS[fill_key]

                    for col in range(1, len(headers) + 1):
                        cell = ws.cell(row=row_idx, column=col)
                        cell.border = EXCEL_BORDER
                        cell.alignment = EXCEL_ALIGNMENT_LEFT
                        cell.font = EXCEL_FONTS["normal"]

                    row_idx += 1

            DocGenerator._auto_adjust_columns(ws)
            DocGenerator._freeze_panes(ws)
            DocGenerator._add_autofilter(ws)

            gap_excel_path = output_dir / "gap_анализ.xlsx"
            wb.save(str(gap_excel_path))
            size_kb = gap_excel_path.stat().st_size / 1024
            logger.info("  [✓] GAP-анализ Excel: %s (%.1f KB)", gap_excel_path.name, size_kb)
        except Exception as exc:
            logger.error("  [✗] Ошибка Excel (GAP): %s", exc)

    # =================================================================
    # Итог
    # =================================================================
    logger.info("=" * 60)
    logger.info("ИТОГ: Генерация документов завершена")
    logger.info("=" * 60)

    # Подсчитываем файлы
    bpmn_files = list(bpmn_dir.glob("*.bpmn"))
    svg_files = list(bpmn_dir.glob("*.svg"))
    vsdx_files = list(visio_dir.glob("*.vsdx"))
    docx_files = list(output_dir.glob("*.docx"))
    xlsx_files = list(output_dir.glob("*.xlsx"))

    logger.info("  BPMN XML:  %d файлов", len(bpmn_files))
    logger.info("  SVG:       %d файлов", len(svg_files))
    logger.info("  Visio:     %d файлов", len(vsdx_files))
    logger.info("  Word:      %d файлов", len(docx_files))
    logger.info("  Excel:     %d файлов", len(xlsx_files))
    logger.info("")
    logger.info("Все файлы: %s", output_dir)


if __name__ == "__main__":
    main()
