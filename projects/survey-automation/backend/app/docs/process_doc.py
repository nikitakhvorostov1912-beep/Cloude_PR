"""Генератор Word-документа с описанием бизнес-процессов.

Создаёт профессионально оформленный документ формата .docx,
содержащий: титульную страницу, оглавление, карточки процессов,
таблицы шагов, решений, проблемных зон и сводку.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .doc_generator import (
    COLORS,
    DocGenerator,
    format_date_russian,
    safe_list,
    safe_str,
)

logger = logging.getLogger(__name__)

# Цвета для уровней критичности проблем (HEX без #)
_SEVERITY_COLORS: dict[str, str] = {
    "critical": "CC0000",
    "high": "FF6600",
    "medium": "FFCC00",
    "low": "339933",
}

_SEVERITY_TEXT_COLORS: dict[str, RGBColor] = {
    "critical": RGBColor(0xFF, 0xFF, 0xFF),
    "high": RGBColor(0xFF, 0xFF, 0xFF),
    "medium": RGBColor(0x33, 0x33, 0x33),
    "low": RGBColor(0xFF, 0xFF, 0xFF),
}

_SEVERITY_LABELS: dict[str, str] = {
    "critical": "Критический",
    "high": "Высокий",
    "medium": "Средний",
    "low": "Низкий",
}


class ProcessDocGenerator(DocGenerator):
    """Генератор Word-документа «Описание бизнес-процессов».

    Создаёт структурированный документ с полным описанием
    бизнес-процессов компании, включая карточки процессов,
    пошаговые описания, точки принятия решений и проблемные зоны.

    Example::

        gen = ProcessDocGenerator(author="Аналитик", company="ООО Рога и Копыта")
        path = gen.generate(
            processes=[{...}, {...}],
            project_name="Автоматизация склада",
            output_path=Path("output/processes.docx"),
        )
    """

    def generate(
        self,
        processes: list[dict[str, Any]],
        output_path: Path,
        project_name: str = "Проект",
    ) -> Path:
        """Генерирует Word-документ с описанием бизнес-процессов.

        Args:
            processes: Список словарей с данными процессов. Каждый процесс
                может содержать поля: name, trigger, result, participants,
                department, steps, decisions, pain_points, description.
            output_path: Путь для сохранения файла .docx.
            project_name: Название проекта (отображается на титуле).

        Returns:
            Путь к созданному документу.
        """
        self._ensure_parent_dir(output_path)

        date_str = format_date_russian()
        doc = self._create_document(title=f"Описание бизнес-процессов — {project_name}")

        # Титульная страница
        self._add_title_page(doc, "Описание бизнес-процессов", project_name, date_str)

        # Оглавление
        self._add_table_of_contents(doc)

        # Верхний колонтитул и нумерация страниц
        self._add_header(doc, f"{project_name} — Описание бизнес-процессов")
        self._add_footer_with_pages(doc)

        # Процессы
        if processes:
            for idx, process in enumerate(processes, start=1):
                self._add_process_section(doc, process, idx)
        else:
            p = doc.add_paragraph("Данные о бизнес-процессах отсутствуют.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Сводка
        self._add_summary_section(doc, processes)

        doc.save(str(output_path))
        logger.info("Документ процессов сохранён: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Секция процесса
    # ------------------------------------------------------------------

    def _add_process_section(
        self,
        doc: Any,
        process: dict[str, Any],
        index: int,
    ) -> None:
        """Добавляет раздел одного бизнес-процесса.

        Args:
            doc: Word-документ.
            process: Данные процесса.
            index: Порядковый номер процесса.
        """
        name = safe_str(process.get("name"), f"Процесс {index}")
        doc.add_heading(f"{index}. {name}", level=1)

        # Карточка процесса
        self._add_process_card(doc, process)

        # Шаги
        steps = process.get("steps", [])
        if steps:
            self._add_steps_table(doc, steps)

        # Решения
        decisions = process.get("decisions", [])
        if decisions:
            self._add_decisions_table(doc, decisions)

        # Проблемные зоны
        pain_points = process.get("pain_points", [])
        if pain_points:
            self._add_pain_points(doc, pain_points)

        # Текстовое описание
        description = process.get("description", "")
        if description:
            self._add_description(doc, description)

        # Разрыв страницы после каждого процесса (кроме последнего)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Карточка процесса (таблица)
    # ------------------------------------------------------------------

    def _add_process_card(self, doc: Any, process: dict[str, Any]) -> None:
        """Добавляет карточку процесса в формате двуколоночной таблицы.

        Поля: Название, Триггер, Результат, Участники, Отдел.

        Args:
            doc: Word-документ.
            process: Данные процесса.
        """
        doc.add_heading("Карточка процесса", level=2)

        card_fields = [
            ("Название", safe_str(process.get("name"))),
            ("Триггер", safe_str(process.get("trigger"))),
            ("Результат", safe_str(process.get("result"))),
            ("Участники", safe_list(process.get("participants"))),
            ("Отдел", safe_str(process.get("department"))),
        ]

        table = doc.add_table(rows=len(card_fields), cols=2)
        table.style = "Table Grid"

        # Ширина столбцов
        for row_data, (label, value) in zip(table.rows, card_fields):
            # Первый столбец — метка
            cell_label = row_data.cells[0]
            cell_label.text = ""
            p = cell_label.paragraphs[0]
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            cell_label.width = Cm(4)
            self._set_cell_shading(cell_label, "E8EDF5")

            # Второй столбец — значение
            cell_value = row_data.cells[1]
            cell_value.text = ""
            p_val = cell_value.paragraphs[0]
            run_val = p_val.add_run(value)
            run_val.font.size = Pt(10)
            run_val.font.name = "Calibri"
            cell_value.width = Cm(13)

        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Таблица шагов
    # ------------------------------------------------------------------

    def _add_steps_table(self, doc: Any, steps: list[dict[str, Any]]) -> None:
        """Добавляет таблицу шагов процесса.

        Столбцы: №, Шаг, Исполнитель, Входы, Выходы, Системы.

        Args:
            doc: Word-документ.
            steps: Список шагов процесса.
        """
        doc.add_heading("Шаги процесса", level=2)

        headers = ["№", "Шаг", "Исполнитель", "Входы", "Выходы", "Системы"]
        table = doc.add_table(rows=1 + len(steps), cols=len(headers))
        table.style = "Table Grid"

        # Заголовок
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._format_table_header(table)

        # Данные
        for row_idx, step in enumerate(steps, start=1):
            row = table.rows[row_idx]
            values = [
                str(step.get("number", row_idx)),
                safe_str(step.get("name", step.get("step", ""))),
                safe_str(step.get("executor", step.get("performer", ""))),
                safe_list(step.get("inputs", step.get("input", []))),
                safe_list(step.get("outputs", step.get("output", []))),
                safe_list(step.get("systems", step.get("system", []))),
            ]
            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._format_table_rows(table)
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Таблица решений
    # ------------------------------------------------------------------

    def _add_decisions_table(self, doc: Any, decisions: list[dict[str, Any]]) -> None:
        """Добавляет таблицу точек принятия решений.

        Столбцы: Условие, Да, Нет.

        Args:
            doc: Word-документ.
            decisions: Список решений.
        """
        doc.add_heading("Точки принятия решений", level=2)

        headers = ["Условие", "Да", "Нет"]
        table = doc.add_table(rows=1 + len(decisions), cols=len(headers))
        table.style = "Table Grid"

        # Заголовок
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._format_table_header(table)

        # Данные
        for row_idx, decision in enumerate(decisions, start=1):
            row = table.rows[row_idx]
            values = [
                safe_str(decision.get("condition", decision.get("question", ""))),
                safe_str(decision.get("yes", decision.get("true_branch", ""))),
                safe_str(decision.get("no", decision.get("false_branch", ""))),
            ]
            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"

        self._format_table_rows(table)
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Проблемные зоны
    # ------------------------------------------------------------------

    def _add_pain_points(self, doc: Any, pain_points: list[dict[str, Any]]) -> None:
        """Добавляет раздел с проблемными зонами, окрашенными по критичности.

        Args:
            doc: Word-документ.
            pain_points: Список проблемных зон с полями severity, description.
        """
        doc.add_heading("Проблемные зоны", level=2)

        for pp in pain_points:
            severity = str(pp.get("severity", "medium")).lower().strip()
            description = safe_str(pp.get("description", pp.get("text", "")))
            label = _SEVERITY_LABELS.get(severity, severity.capitalize())
            color_hex = _SEVERITY_COLORS.get(severity, "FFCC00")
            text_color = _SEVERITY_TEXT_COLORS.get(severity, RGBColor(0x33, 0x33, 0x33))

            # Каждая проблема — маленькая таблица из одной строки для фона
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            cell.text = ""
            self._set_cell_shading(cell, color_hex)

            p = cell.paragraphs[0]
            # Метка критичности
            run_label = p.add_run(f"[{label}] ")
            run_label.font.bold = True
            run_label.font.size = Pt(10)
            run_label.font.name = "Calibri"
            run_label.font.color.rgb = text_color

            # Описание
            run_desc = p.add_run(description)
            run_desc.font.size = Pt(10)
            run_desc.font.name = "Calibri"
            run_desc.font.color.rgb = text_color

            # Рекомендация (если есть)
            recommendation = pp.get("recommendation", "")
            if recommendation:
                p_rec = cell.add_paragraph()
                run_rec_label = p_rec.add_run("Рекомендация: ")
                run_rec_label.font.bold = True
                run_rec_label.font.size = Pt(9)
                run_rec_label.font.name = "Calibri"
                run_rec_label.font.color.rgb = text_color
                run_rec = p_rec.add_run(str(recommendation))
                run_rec.font.size = Pt(9)
                run_rec.font.name = "Calibri"
                run_rec.font.color.rgb = text_color

        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Текстовое описание
    # ------------------------------------------------------------------

    def _add_description(self, doc: Any, description: str) -> None:
        """Добавляет текстовое описание процесса.

        Args:
            doc: Word-документ.
            description: Свободный текст описания.
        """
        doc.add_heading("Описание", level=2)

        paragraphs = str(description).strip().split("\n")
        for para_text in paragraphs:
            text = para_text.strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)

    # ------------------------------------------------------------------
    # Сводка
    # ------------------------------------------------------------------

    def _add_summary_section(self, doc: Any, processes: list[dict[str, Any]]) -> None:
        """Добавляет итоговую сводку по всем процессам.

        Содержит общую статистику: количество процессов, шагов,
        решений и проблемных зон.

        Args:
            doc: Word-документ.
            processes: Список всех процессов.
        """
        doc.add_heading("Сводка", level=1)

        total_processes = len(processes)
        total_steps = sum(len(p.get("steps", [])) for p in processes)
        total_decisions = sum(len(p.get("decisions", [])) for p in processes)
        total_pain_points = sum(len(p.get("pain_points", [])) for p in processes)

        # Подсчёт по критичности
        severity_counts: dict[str, int] = {}
        for proc in processes:
            for pp in proc.get("pain_points", []):
                sev = str(pp.get("severity", "medium")).lower().strip()
                severity_counts[sev] = severity_counts.get(sev, 0) + 1

        # Подсчёт по отделам
        departments: dict[str, int] = {}
        for proc in processes:
            dept = safe_str(proc.get("department"), "Не указан")
            departments[dept] = departments.get(dept, 0) + 1

        # Таблица статистики
        doc.add_heading("Общая статистика", level=2)

        stats = [
            ("Всего процессов", str(total_processes)),
            ("Всего шагов", str(total_steps)),
            ("Точек принятия решений", str(total_decisions)),
            ("Проблемных зон", str(total_pain_points)),
        ]

        table_stats = doc.add_table(rows=len(stats), cols=2)
        table_stats.style = "Table Grid"

        for row_data, (label, value) in zip(table_stats.rows, stats):
            cell_label = row_data.cells[0]
            cell_label.text = ""
            p_l = cell_label.paragraphs[0]
            run_l = p_l.add_run(label)
            run_l.font.bold = True
            run_l.font.size = Pt(10)
            run_l.font.name = "Calibri"
            self._set_cell_shading(cell_label, "E8EDF5")

            cell_value = row_data.cells[1]
            cell_value.text = ""
            p_v = cell_value.paragraphs[0]
            run_v = p_v.add_run(value)
            run_v.font.size = Pt(10)
            run_v.font.name = "Calibri"
            p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Проблемы по критичности
        if severity_counts:
            doc.add_heading("Проблемные зоны по критичности", level=2)

            table_sev = doc.add_table(rows=1 + len(severity_counts), cols=2)
            table_sev.style = "Table Grid"

            # Заголовок
            table_sev.rows[0].cells[0].text = ""
            p0 = table_sev.rows[0].cells[0].paragraphs[0]
            r0 = p0.add_run("Критичность")
            r0.font.bold = True
            r0.font.size = Pt(9)
            r0.font.name = "Calibri"

            table_sev.rows[0].cells[1].text = ""
            p1 = table_sev.rows[0].cells[1].paragraphs[0]
            r1 = p1.add_run("Количество")
            r1.font.bold = True
            r1.font.size = Pt(9)
            r1.font.name = "Calibri"
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._format_table_header(table_sev)

            # Порядок: critical, high, medium, low
            order = ["critical", "high", "medium", "low"]
            sorted_items = sorted(
                severity_counts.items(),
                key=lambda x: order.index(x[0]) if x[0] in order else 99,
            )

            for row_idx, (sev, count) in enumerate(sorted_items, start=1):
                row = table_sev.rows[row_idx]
                label = _SEVERITY_LABELS.get(sev, sev.capitalize())
                color_hex = _SEVERITY_COLORS.get(sev, "FFCC00")

                cell_sev = row.cells[0]
                cell_sev.text = ""
                p_s = cell_sev.paragraphs[0]
                run_s = p_s.add_run(label)
                run_s.font.size = Pt(9)
                run_s.font.name = "Calibri"
                self._set_cell_shading(cell_sev, color_hex)

                text_color = _SEVERITY_TEXT_COLORS.get(sev, RGBColor(0x33, 0x33, 0x33))
                run_s.font.color.rgb = text_color

                cell_cnt = row.cells[1]
                cell_cnt.text = ""
                p_c = cell_cnt.paragraphs[0]
                run_c = p_c.add_run(str(count))
                run_c.font.size = Pt(9)
                run_c.font.name = "Calibri"
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")

        # Процессы по отделам
        if departments and len(departments) > 1:
            doc.add_heading("Процессы по отделам", level=2)

            table_dept = doc.add_table(rows=1 + len(departments), cols=2)
            table_dept.style = "Table Grid"

            table_dept.rows[0].cells[0].text = ""
            pd0 = table_dept.rows[0].cells[0].paragraphs[0]
            rd0 = pd0.add_run("Отдел")
            rd0.font.bold = True
            rd0.font.size = Pt(9)
            rd0.font.name = "Calibri"

            table_dept.rows[0].cells[1].text = ""
            pd1 = table_dept.rows[0].cells[1].paragraphs[0]
            rd1 = pd1.add_run("Процессов")
            rd1.font.bold = True
            rd1.font.size = Pt(9)
            rd1.font.name = "Calibri"
            pd1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._format_table_header(table_dept)

            for row_idx, (dept, count) in enumerate(
                sorted(departments.items(), key=lambda x: x[1], reverse=True),
                start=1,
            ):
                row = table_dept.rows[row_idx]
                row.cells[0].text = ""
                pd = row.cells[0].paragraphs[0]
                rd = pd.add_run(dept)
                rd.font.size = Pt(9)
                rd.font.name = "Calibri"

                row.cells[1].text = ""
                pc = row.cells[1].paragraphs[0]
                rc = pc.add_run(str(count))
                rc.font.size = Pt(9)
                rc.font.name = "Calibri"
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._format_table_rows(table_dept)
