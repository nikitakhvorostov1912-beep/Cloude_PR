"""Генерация отчёта GAP-анализа (Excel)."""

from __future__ import annotations

import logging
from collections import Counter
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill

from ..exceptions import ExportError
from .doc_generator import (
    COLORS,
    EXCEL_ALIGNMENT_CENTER,
    EXCEL_ALIGNMENT_LEFT,
    EXCEL_BORDER,
    EXCEL_FILLS,
    EXCEL_FONTS,
    DocGenerator,
    format_date_russian,
    safe_str,
)

logger = logging.getLogger(__name__)

# Заголовки основной таблицы GAP-анализа
_GAP_HEADERS: list[str] = [
    "Процесс",
    "Функция",
    "Покрытие (%)",
    "Модуль ERP",
    "Описание GAP",
    "Рекомендация",
    "Трудоёмкость (чел/дни)",
]

# Пороги покрытия и соответствующие цвета
_COVERAGE_FILLS: dict[str, PatternFill] = {
    "green": PatternFill(start_color="FF339933", end_color="FF339933", fill_type="solid"),
    "yellow": PatternFill(start_color="FFFFCC00", end_color="FFFFCC00", fill_type="solid"),
    "red": PatternFill(start_color="FFCC0000", end_color="FFCC0000", fill_type="solid"),
}

_COVERAGE_FONT_COLORS: dict[str, str] = {
    "green": "FFFFFFFF",
    "yellow": "FF333333",
    "red": "FFFFFFFF",
}

_COVERAGE_HEX_WORD: dict[str, str] = {
    "green": "339933",
    "yellow": "FFCC00",
    "red": "CC0000",
}

_COVERAGE_TEXT_WORD: dict[str, RGBColor] = {
    "green": RGBColor(0xFF, 0xFF, 0xFF),
    "yellow": RGBColor(0x33, 0x33, 0x33),
    "red": RGBColor(0xFF, 0xFF, 0xFF),
}


def _coverage_level(value: float) -> str:
    """Определяет уровень покрытия по значению процента.

    Args:
        value: Значение покрытия (0-100).

    Returns:
        Ключ уровня: «green», «yellow» или «red».
    """
    if value >= 80:
        return "green"
    if value >= 50:
        return "yellow"
    return "red"


class GapReportGenerator(DocGenerator):
    """Генератор отчётов GAP-анализа в форматах Excel и Word.

    Создаёт отчёт с анализом функциональных разрывов между текущими
    бизнес-процессами и возможностями ERP-системы.

    Example::

        gen = GapReportGenerator(author="Аналитик", company="ООО Компания")
        gen.generate_excel(
            gaps=[{...}],
            project_name="ERP-внедрение",
            erp_config="1С:ERP 2.5",
            output_path=Path("output/gap_analysis.xlsx"),
        )
    """

    def generate(
        self,
        gaps: list[dict[str, Any]],
        output_path: Path,
        project_name: str = "Проект",
        erp_config: str = "",
    ) -> Path:
        """Генерирует отчёт GAP-анализа (Excel по умолчанию).

        Args:
            gaps: Список GAP-записей.
            output_path: Путь для сохранения.
            project_name: Название проекта.
            erp_config: Конфигурация ERP-системы.

        Returns:
            Путь к созданному документу.
        """
        if output_path.suffix.lower() == ".docx":
            return self.generate_word(gaps, project_name, erp_config, output_path)
        return self.generate_excel(gaps, project_name, erp_config, output_path)

    # ------------------------------------------------------------------
    # Excel-отчёт
    # ------------------------------------------------------------------

    def generate_excel(
        self,
        gaps: list[dict[str, Any]],
        project_name: str,
        erp_config: str,
        output_path: Path,
    ) -> Path:
        """Генерирует Excel-отчёт GAP-анализа.

        Args:
            gaps: Список словарей с полями: process, function, coverage,
                erp_module, gap_description, recommendation, effort.
            project_name: Название проекта.
            erp_config: Конфигурация ERP-системы.
            output_path: Путь для сохранения .xlsx файла.

        Returns:
            Путь к созданному документу.

        Raises:
            ExportError: При ошибке генерации отчёта.
        """
        try:
            self._ensure_parent_dir(output_path)
            wb = self._create_workbook()

            # Лист «GAP-анализ»
            ws_gap = wb.active
            ws_gap.title = "GAP-анализ"
            self._fill_gap_sheet(ws_gap, gaps, project_name, erp_config)

            # Лист «Сводка по модулям»
            ws_summary = wb.create_sheet("Сводка по модулям")
            self._fill_module_summary_sheet(ws_summary, gaps, project_name)

            # Лист «Рекомендации»
            ws_recs = wb.create_sheet("Рекомендации")
            self._fill_recommendations_sheet(ws_recs, gaps, project_name)

            wb.save(str(output_path))
            logger.info("Отчёт GAP-анализа (Excel) сохранён: %s", output_path)
            return output_path

        except Exception as exc:
            logger.exception("Ошибка генерации Excel-отчёта GAP-анализа: %s", exc)
            raise ExportError(
                "Ошибка при генерации отчёта GAP-анализа (Excel)",
                detail=str(exc),
            ) from exc

    def _fill_gap_sheet(
        self,
        ws: Any,
        gaps: list[dict[str, Any]],
        project_name: str,
        erp_config: str,
    ) -> None:
        """Заполняет основной лист GAP-анализа.

        Args:
            ws: Рабочий лист Excel.
            gaps: Список GAP-записей.
            project_name: Название проекта.
            erp_config: Конфигурация ERP.
        """
        # Заголовок
        ws.merge_cells(
            start_row=1, start_column=1,
            end_row=1, end_column=len(_GAP_HEADERS),
        )
        title_text = f"GAP-анализ — {project_name}"
        if erp_config:
            title_text += f" ({erp_config})"
        title_cell = ws.cell(row=1, column=1, value=title_text)
        title_cell.font = EXCEL_FONTS["title"]
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        # Дата
        ws.merge_cells(
            start_row=2, start_column=1,
            end_row=2, end_column=len(_GAP_HEADERS),
        )
        date_cell = ws.cell(row=2, column=1, value=format_date_russian())
        date_cell.font = Font(name="Calibri", size=10, italic=True, color="FF999999")
        date_cell.alignment = Alignment(horizontal="center")

        # Заголовки таблицы (строка 4)
        header_row = 4
        self._write_excel_header(ws, _GAP_HEADERS, row=header_row)

        # Данные
        for row_idx, gap in enumerate(gaps, start=header_row + 1):
            coverage_raw = gap.get("coverage", 0)
            try:
                coverage_val = float(coverage_raw)
            except (ValueError, TypeError):
                coverage_val = 0.0

            values: list[Any] = [
                safe_str(gap.get("process")),
                safe_str(gap.get("function")),
                round(coverage_val, 1),
                safe_str(gap.get("erp_module")),
                safe_str(gap.get("gap_description")),
                safe_str(gap.get("recommendation")),
                safe_str(gap.get("effort")),
            ]

            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = EXCEL_FONTS["normal"]
                cell.border = EXCEL_BORDER

                if col_idx in (3, 7):
                    cell.alignment = EXCEL_ALIGNMENT_CENTER
                else:
                    cell.alignment = EXCEL_ALIGNMENT_LEFT

                # Чередование строк
                if (row_idx - header_row) % 2 == 0:
                    cell.fill = EXCEL_FILLS["alt_row"]

            # Условное форматирование покрытия (столбец 3)
            coverage_cell = ws.cell(row=row_idx, column=3)
            level = _coverage_level(coverage_val)
            coverage_cell.fill = _COVERAGE_FILLS[level]
            coverage_cell.font = Font(
                name="Calibri",
                size=11,
                bold=True,
                color=_COVERAGE_FONT_COLORS[level],
            )

        # Автофильтр, закрепление, ширины
        self._add_autofilter(ws, row=header_row)
        self._freeze_panes(ws, f"A{header_row + 1}")
        self._auto_adjust_columns(ws)

        # Принудительные ширины
        ws.column_dimensions["A"].width = 25  # Процесс
        ws.column_dimensions["B"].width = 25  # Функция
        ws.column_dimensions["E"].width = 40  # Описание GAP
        ws.column_dimensions["F"].width = 40  # Рекомендация

    def _fill_module_summary_sheet(
        self,
        ws: Any,
        gaps: list[dict[str, Any]],
        project_name: str,
    ) -> None:
        """Заполняет лист сводки по ERP-модулям.

        Args:
            ws: Рабочий лист Excel.
            gaps: Список GAP-записей.
            project_name: Название проекта.
        """
        # Заголовок
        summary_headers = [
            "Модуль ERP",
            "Кол-во функций",
            "Среднее покрытие (%)",
            "Мин. покрытие (%)",
            "Макс. покрытие (%)",
            "Кол-во GAP",
        ]

        ws.merge_cells(
            start_row=1, start_column=1,
            end_row=1, end_column=len(summary_headers),
        )
        title_cell = ws.cell(
            row=1, column=1,
            value=f"Сводка по модулям ERP — {project_name}",
        )
        title_cell.font = EXCEL_FONTS["title"]
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        header_row = 3
        self._write_excel_header(ws, summary_headers, row=header_row)

        # Агрегация по модулям
        modules: dict[str, list[float]] = {}
        gap_counts: dict[str, int] = {}
        for gap in gaps:
            module = safe_str(gap.get("erp_module"), "Не указан")
            try:
                coverage_val = float(gap.get("coverage", 0))
            except (ValueError, TypeError):
                coverage_val = 0.0

            if module not in modules:
                modules[module] = []
                gap_counts[module] = 0
            modules[module].append(coverage_val)

            gap_desc = safe_str(gap.get("gap_description"), "")
            if gap_desc and gap_desc != "—":
                gap_counts[module] += 1

        # Данные
        sorted_modules = sorted(modules.items(), key=lambda x: sum(x[1]) / len(x[1]) if x[1] else 0)
        for row_idx, (module_name, coverages) in enumerate(sorted_modules, start=header_row + 1):
            avg_cov = sum(coverages) / len(coverages) if coverages else 0.0
            min_cov = min(coverages) if coverages else 0.0
            max_cov = max(coverages) if coverages else 0.0
            num_gaps = gap_counts.get(module_name, 0)

            values: list[Any] = [
                module_name,
                len(coverages),
                round(avg_cov, 1),
                round(min_cov, 1),
                round(max_cov, 1),
                num_gaps,
            ]

            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = EXCEL_FONTS["normal"]
                cell.border = EXCEL_BORDER

                if col_idx >= 2:
                    cell.alignment = EXCEL_ALIGNMENT_CENTER
                else:
                    cell.alignment = EXCEL_ALIGNMENT_LEFT

                # Чередование строк
                if (row_idx - header_row) % 2 == 0:
                    cell.fill = EXCEL_FILLS["alt_row"]

            # Цвет среднего покрытия
            avg_cell = ws.cell(row=row_idx, column=3)
            level = _coverage_level(avg_cov)
            avg_cell.fill = _COVERAGE_FILLS[level]
            avg_cell.font = Font(
                name="Calibri",
                size=11,
                bold=True,
                color=_COVERAGE_FONT_COLORS[level],
            )

        # Автофильтр, закрепление, ширины
        self._add_autofilter(ws, row=header_row)
        self._freeze_panes(ws, f"A{header_row + 1}")
        self._auto_adjust_columns(ws)
        ws.column_dimensions["A"].width = 30

    def _fill_recommendations_sheet(
        self,
        ws: Any,
        gaps: list[dict[str, Any]],
        project_name: str,
    ) -> None:
        """Заполняет лист рекомендаций (приоритезированный список доработок).

        Args:
            ws: Рабочий лист Excel.
            gaps: Список GAP-записей.
            project_name: Название проекта.
        """
        rec_headers = [
            "№",
            "Процесс",
            "Функция",
            "Покрытие (%)",
            "Рекомендация",
            "Модуль ERP",
            "Трудоёмкость (чел/дни)",
        ]

        ws.merge_cells(
            start_row=1, start_column=1,
            end_row=1, end_column=len(rec_headers),
        )
        title_cell = ws.cell(
            row=1, column=1,
            value=f"Рекомендации по доработкам — {project_name}",
        )
        title_cell.font = EXCEL_FONTS["title"]
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        header_row = 3
        self._write_excel_header(ws, rec_headers, row=header_row)

        # Фильтруем записи с рекомендациями и сортируем по покрытию (по возрастанию)
        recs_with_data = []
        for gap in gaps:
            rec_text = safe_str(gap.get("recommendation"), "")
            if rec_text and rec_text != "—":
                try:
                    coverage_val = float(gap.get("coverage", 0))
                except (ValueError, TypeError):
                    coverage_val = 0.0
                recs_with_data.append((gap, coverage_val))

        recs_with_data.sort(key=lambda x: x[1])

        for num, (gap, coverage_val) in enumerate(recs_with_data, start=1):
            row_idx = header_row + num
            values: list[Any] = [
                num,
                safe_str(gap.get("process")),
                safe_str(gap.get("function")),
                round(coverage_val, 1),
                safe_str(gap.get("recommendation")),
                safe_str(gap.get("erp_module")),
                safe_str(gap.get("effort")),
            ]

            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = EXCEL_FONTS["normal"]
                cell.border = EXCEL_BORDER

                if col_idx in (1, 4, 7):
                    cell.alignment = EXCEL_ALIGNMENT_CENTER
                else:
                    cell.alignment = EXCEL_ALIGNMENT_LEFT

                # Чередование строк
                if num % 2 == 0:
                    cell.fill = EXCEL_FILLS["alt_row"]

            # Цвет покрытия
            cov_cell = ws.cell(row=row_idx, column=4)
            level = _coverage_level(coverage_val)
            cov_cell.fill = _COVERAGE_FILLS[level]
            cov_cell.font = Font(
                name="Calibri",
                size=11,
                bold=True,
                color=_COVERAGE_FONT_COLORS[level],
            )

        # Автофильтр, закрепление, ширины
        self._add_autofilter(ws, row=header_row)
        self._freeze_panes(ws, f"A{header_row + 1}")
        self._auto_adjust_columns(ws)
        ws.column_dimensions["B"].width = 25
        ws.column_dimensions["C"].width = 25
        ws.column_dimensions["E"].width = 50

    # ------------------------------------------------------------------
    # Word-отчёт
    # ------------------------------------------------------------------

    def generate_word(
        self,
        gaps: list[dict[str, Any]],
        project_name: str,
        erp_config: str,
        output_path: Path,
    ) -> Path:
        """Генерирует Word-отчёт GAP-анализа.

        Args:
            gaps: Список GAP-записей.
            project_name: Название проекта.
            erp_config: Конфигурация ERP-системы.
            output_path: Путь для сохранения .docx файла.

        Returns:
            Путь к созданному документу.

        Raises:
            ExportError: При ошибке генерации отчёта.
        """
        try:
            self._ensure_parent_dir(output_path)
            date_str = format_date_russian()
            doc = self._create_document(title=f"GAP-анализ — {project_name}")

            # Титульная страница
            title = "Отчёт GAP-анализа"
            if erp_config:
                title += f"\n{erp_config}"
            self._add_title_page(doc, title, project_name, date_str)

            # Оглавление
            self._add_table_of_contents(doc)

            # Колонтитулы
            self._add_header(doc, f"{project_name} — GAP-анализ")
            self._add_footer_with_pages(doc)

            # Резюме
            self._add_word_executive_summary(doc, gaps, erp_config)

            # Детальная таблица GAP
            self._add_word_gap_table(doc, gaps)

            # Рекомендации
            self._add_word_recommendations(doc, gaps)

            doc.save(str(output_path))
            logger.info("Отчёт GAP-анализа (Word) сохранён: %s", output_path)
            return output_path

        except Exception as exc:
            logger.exception("Ошибка генерации Word-отчёта GAP-анализа: %s", exc)
            raise ExportError(
                "Ошибка при генерации отчёта GAP-анализа (Word)",
                detail=str(exc),
            ) from exc

    def _add_word_executive_summary(
        self,
        doc: Any,
        gaps: list[dict[str, Any]],
        erp_config: str,
    ) -> None:
        """Добавляет раздел «Резюме» в Word-отчёт.

        Args:
            doc: Word-документ.
            gaps: Список GAP-записей.
            erp_config: Конфигурация ERP.
        """
        doc.add_heading("Резюме", level=1)

        # Подсчёт статистики
        total_functions = len(gaps)
        coverages = []
        for gap in gaps:
            try:
                coverages.append(float(gap.get("coverage", 0)))
            except (ValueError, TypeError):
                coverages.append(0.0)

        avg_coverage = sum(coverages) / len(coverages) if coverages else 0.0
        green_count = sum(1 for c in coverages if c >= 80)
        yellow_count = sum(1 for c in coverages if 50 <= c < 80)
        red_count = sum(1 for c in coverages if c < 50)

        # Описание
        if erp_config:
            doc.add_paragraph(
                f"Отчёт содержит результаты GAP-анализа бизнес-процессов "
                f"относительно конфигурации «{erp_config}»."
            )

        # Таблица статистики
        doc.add_heading("Общая статистика", level=2)
        stats = [
            ("Всего функций проанализировано", str(total_functions)),
            ("Среднее покрытие", f"{avg_coverage:.1f}%"),
            ("Полное покрытие (80-100%)", str(green_count)),
            ("Частичное покрытие (50-79%)", str(yellow_count)),
            ("Низкое покрытие (0-49%)", str(red_count)),
        ]

        table_stats = doc.add_table(rows=len(stats), cols=2)
        table_stats.style = "Table Grid"

        for row_data, (label, value) in zip(table_stats.rows, stats):
            cell_label = row_data.cells[0]
            cell_label.text = ""
            p_l = cell_label.paragraphs[0]
            run_l = p_l.add_run(label)
            run_l.font.bold = True
            run_l.font.size = Pt(10)
            run_l.font.name = "Calibri"
            self._set_cell_shading(cell_label, "E8EDF5")

            cell_value = row_data.cells[1]
            cell_value.text = ""
            p_v = cell_value.paragraphs[0]
            run_v = p_v.add_run(value)
            run_v.font.size = Pt(10)
            run_v.font.name = "Calibri"
            p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        doc.add_page_break()

    def _add_word_gap_table(
        self,
        doc: Any,
        gaps: list[dict[str, Any]],
    ) -> None:
        """Добавляет детальную таблицу GAP-анализа в Word.

        Args:
            doc: Word-документ.
            gaps: Список GAP-записей.
        """
        doc.add_heading("Детальный GAP-анализ", level=1)

        if not gaps:
            p = doc.add_paragraph("Данные GAP-анализа отсутствуют.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        headers = _GAP_HEADERS
        table = doc.add_table(rows=1 + len(gaps), cols=len(headers))
        table.style = "Table Grid"

        # Заголовок
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._format_table_header(table)

        # Данные
        for row_idx, gap in enumerate(gaps, start=1):
            row = table.rows[row_idx]

            try:
                coverage_val = float(gap.get("coverage", 0))
            except (ValueError, TypeError):
                coverage_val = 0.0

            values = [
                safe_str(gap.get("process")),
                safe_str(gap.get("function")),
                f"{coverage_val:.0f}%",
                safe_str(gap.get("erp_module")),
                safe_str(gap.get("gap_description")),
                safe_str(gap.get("recommendation")),
                safe_str(gap.get("effort")),
            ]

            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(8)
                run.font.name = "Calibri"
                if col_idx in (2, 6):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Цвет ячейки покрытия
            coverage_cell = row.cells[2]
            level = _coverage_level(coverage_val)
            self._set_cell_shading(coverage_cell, _COVERAGE_HEX_WORD[level])
            text_color = _COVERAGE_TEXT_WORD[level]
            for p in coverage_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = text_color

        self._format_table_rows(table)
        doc.add_page_break()

    def _add_word_recommendations(
        self,
        doc: Any,
        gaps: list[dict[str, Any]],
    ) -> None:
        """Добавляет раздел рекомендаций в Word-отчёт.

        Args:
            doc: Word-документ.
            gaps: Список GAP-записей.
        """
        doc.add_heading("Рекомендации", level=1)

        # Фильтруем записи с рекомендациями, сортируем по покрытию
        recs = []
        for gap in gaps:
            rec_text = safe_str(gap.get("recommendation"), "")
            if rec_text and rec_text != "—":
                try:
                    cov = float(gap.get("coverage", 0))
                except (ValueError, TypeError):
                    cov = 0.0
                recs.append((gap, cov))

        recs.sort(key=lambda x: x[1])

        if not recs:
            p = doc.add_paragraph("Рекомендации отсутствуют.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        for num, (gap, cov) in enumerate(recs, start=1):
            process = safe_str(gap.get("process"))
            function = safe_str(gap.get("function"))
            rec_text = safe_str(gap.get("recommendation"))
            effort = safe_str(gap.get("effort"))
            erp_module = safe_str(gap.get("erp_module"))

            doc.add_heading(
                f"{num}. {process} — {function} (покрытие {cov:.0f}%)",
                level=2,
            )

            # Карточка рекомендации
            card_fields = [
                ("Модуль ERP", erp_module),
                ("Рекомендация", rec_text),
                ("Трудоёмкость", effort),
            ]

            table = doc.add_table(rows=len(card_fields), cols=2)
            table.style = "Table Grid"

            for row_data, (label, value) in zip(table.rows, card_fields):
                cell_label = row_data.cells[0]
                cell_label.text = ""
                p_l = cell_label.paragraphs[0]
                run_l = p_l.add_run(label)
                run_l.font.bold = True
                run_l.font.size = Pt(10)
                run_l.font.name = "Calibri"
                run_l.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
                cell_label.width = Cm(4)
                self._set_cell_shading(cell_label, "E8EDF5")

                cell_value = row_data.cells[1]
                cell_value.text = ""
                p_v = cell_value.paragraphs[0]
                run_v = p_v.add_run(value)
                run_v.font.size = Pt(10)
                run_v.font.name = "Calibri"
                cell_value.width = Cm(13)

            doc.add_paragraph("")
