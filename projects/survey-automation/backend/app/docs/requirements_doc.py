"""Генерация документов требований (Word + Excel)."""

from __future__ import annotations

import logging
from collections import Counter
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill

from ..exceptions import ExportError
from .doc_generator import (
    EXCEL_ALIGNMENT_CENTER,
    EXCEL_ALIGNMENT_LEFT,
    EXCEL_BORDER,
    EXCEL_FILLS,
    EXCEL_FONTS,
    DocGenerator,
    format_date_russian,
    safe_str,
)

logger = logging.getLogger(__name__)

# Метки приоритетов MoSCoW
_MOSCOW_LABELS: dict[str, str] = {
    "must": "Must",
    "should": "Should",
    "could": "Could",
    "wont": "Won't",
}

_MOSCOW_FILLS: dict[str, PatternFill] = {
    "must": EXCEL_FILLS["must"],
    "should": EXCEL_FILLS["should"],
    "could": EXCEL_FILLS["could"],
    "wont": EXCEL_FILLS["wont"],
}

_MOSCOW_FONT_COLORS: dict[str, str] = {
    "must": "FFFFFFFF",
    "should": "FFFFFFFF",
    "could": "FF333333",
    "wont": "FFFFFFFF",
}

_MOSCOW_CELL_COLORS_HEX: dict[str, str] = {
    "must": "CC0000",
    "should": "FF6600",
    "could": "FFCC00",
    "wont": "999999",
}

_MOSCOW_TEXT_COLORS_WORD: dict[str, RGBColor] = {
    "must": RGBColor(0xFF, 0xFF, 0xFF),
    "should": RGBColor(0xFF, 0xFF, 0xFF),
    "could": RGBColor(0x33, 0x33, 0x33),
    "wont": RGBColor(0xFF, 0xFF, 0xFF),
}

# Метки типов требований
_TYPE_LABELS: dict[str, str] = {
    "fr": "FR",
    "nfr": "NFR",
    "ir": "IR",
}

# Заголовки таблицы требований
_WORD_HEADERS: list[str] = [
    "ID",
    "Тип",
    "Модуль",
    "Описание",
    "Приоритет",
    "Источник",
    "Трудоёмкость",
]

_EXCEL_HEADERS: list[str] = [
    "ID",
    "Тип (FR/NFR/IR)",
    "Модуль",
    "Описание",
    "Приоритет (MoSCoW)",
    "Источник",
    "Трудоёмкость",
]


class RequirementsDocGenerator(DocGenerator):
    """Генератор документов требований в форматах Word и Excel.

    Создаёт документы со структурированным списком требований,
    включая сводку по приоритетам и типам, цветовое кодирование
    MoSCoW-приоритетов и автофильтры.

    Example::

        gen = RequirementsDocGenerator(author="Аналитик", company="ООО Компания")
        gen.generate_word(
            requirements=[{...}],
            project_name="ERP-внедрение",
            output_path=Path("output/requirements.docx"),
        )
        gen.generate_excel(
            requirements=[{...}],
            project_name="ERP-внедрение",
            output_path=Path("output/requirements.xlsx"),
        )
    """

    def generate(
        self,
        requirements: list[dict[str, Any]],
        output_path: Path,
        project_name: str = "Проект",
    ) -> Path:
        """Генерирует документ требований (Word по умолчанию).

        Args:
            requirements: Список словарей с требованиями.
            output_path: Путь для сохранения файла.
            project_name: Название проекта.

        Returns:
            Путь к созданному документу.
        """
        if output_path.suffix.lower() == ".xlsx":
            return self.generate_excel(requirements, project_name, output_path)
        return self.generate_word(requirements, project_name, output_path)

    # ------------------------------------------------------------------
    # Word-документ
    # ------------------------------------------------------------------

    def generate_word(
        self,
        requirements: list[dict[str, Any]],
        project_name: str,
        output_path: Path,
    ) -> Path:
        """Генерирует Word-документ с таблицей требований.

        Args:
            requirements: Список словарей с полями: id, type, module,
                description, priority, source, effort.
            project_name: Название проекта.
            output_path: Путь для сохранения .docx файла.

        Returns:
            Путь к созданному документу.

        Raises:
            ExportError: При ошибке генерации документа.
        """
        try:
            self._ensure_parent_dir(output_path)
            date_str = format_date_russian()
            doc = self._create_document(title=f"Требования — {project_name}")

            # Титульная страница
            self._add_title_page(doc, "Требования к системе", project_name, date_str)

            # Оглавление
            self._add_table_of_contents(doc)

            # Колонтитулы
            self._add_header(doc, f"{project_name} — Требования к системе")
            self._add_footer_with_pages(doc)

            # Сводка
            self._add_word_summary(doc, requirements)

            # Таблица требований
            self._add_word_requirements_table(doc, requirements)

            doc.save(str(output_path))
            logger.info("Документ требований (Word) сохранён: %s", output_path)
            return output_path

        except Exception as exc:
            logger.exception("Ошибка генерации Word-документа требований: %s", exc)
            raise ExportError(
                "Ошибка при генерации документа требований (Word)",
                detail=str(exc),
            ) from exc

    def _add_word_summary(
        self,
        doc: Any,
        requirements: list[dict[str, Any]],
    ) -> None:
        """Добавляет сводку по требованиям в Word-документ.

        Args:
            doc: Word-документ.
            requirements: Список требований.
        """
        doc.add_heading("Сводка", level=1)

        # Подсчёт по приоритетам
        priority_counts = Counter(
            str(r.get("priority", "")).lower().strip() for r in requirements
        )
        # Подсчёт по типам
        type_counts = Counter(
            str(r.get("type", "")).upper().strip() for r in requirements
        )

        # Таблица: общее количество
        doc.add_heading("Общая статистика", level=2)

        stats = [
            ("Всего требований", str(len(requirements))),
        ]

        # По приоритетам
        moscow_order = ["must", "should", "could", "wont"]
        for key in moscow_order:
            count = priority_counts.get(key, 0)
            label = _MOSCOW_LABELS.get(key, key)
            stats.append((f"Приоритет {label}", str(count)))

        table_stats = doc.add_table(rows=len(stats), cols=2)
        table_stats.style = "Table Grid"

        for row_data, (label, value) in zip(table_stats.rows, stats):
            cell_label = row_data.cells[0]
            cell_label.text = ""
            p_l = cell_label.paragraphs[0]
            run_l = p_l.add_run(label)
            run_l.font.bold = True
            run_l.font.size = Pt(10)
            run_l.font.name = "Calibri"
            self._set_cell_shading(cell_label, "E8EDF5")

            cell_value = row_data.cells[1]
            cell_value.text = ""
            p_v = cell_value.paragraphs[0]
            run_v = p_v.add_run(value)
            run_v.font.size = Pt(10)
            run_v.font.name = "Calibri"
            p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # По типам
        if type_counts:
            doc.add_heading("По типам требований", level=2)

            type_rows = []
            for type_key in ["FR", "NFR", "IR"]:
                count = type_counts.get(type_key, 0)
                if count > 0:
                    type_rows.append((type_key, str(count)))
            # Прочие типы
            for type_key, count in sorted(type_counts.items()):
                if type_key not in {"FR", "NFR", "IR"}:
                    type_rows.append((type_key or "Не указан", str(count)))

            if type_rows:
                table_types = doc.add_table(rows=1 + len(type_rows), cols=2)
                table_types.style = "Table Grid"

                # Заголовок
                for col_idx, header_text in enumerate(["Тип", "Количество"]):
                    cell = table_types.rows[0].cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(header_text)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                self._format_table_header(table_types)

                for row_idx, (type_label, count) in enumerate(type_rows, start=1):
                    row = table_types.rows[row_idx]
                    row.cells[0].text = ""
                    p_t = row.cells[0].paragraphs[0]
                    run_t = p_t.add_run(type_label)
                    run_t.font.size = Pt(9)
                    run_t.font.name = "Calibri"
                    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    row.cells[1].text = ""
                    p_c = row.cells[1].paragraphs[0]
                    run_c = p_c.add_run(count)
                    run_c.font.size = Pt(9)
                    run_c.font.name = "Calibri"
                    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

                self._format_table_rows(table_types)

            doc.add_paragraph("")

        doc.add_page_break()

    def _add_word_requirements_table(
        self,
        doc: Any,
        requirements: list[dict[str, Any]],
    ) -> None:
        """Добавляет основную таблицу требований в Word.

        Args:
            doc: Word-документ.
            requirements: Список требований.
        """
        doc.add_heading("Требования", level=1)

        if not requirements:
            p = doc.add_paragraph("Требования отсутствуют.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        table = doc.add_table(rows=1 + len(requirements), cols=len(_WORD_HEADERS))
        table.style = "Table Grid"

        # Заголовок
        for col_idx, header_text in enumerate(_WORD_HEADERS):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._format_table_header(table)

        # Данные
        for row_idx, req in enumerate(requirements, start=1):
            row = table.rows[row_idx]
            req_type = str(req.get("type", "")).upper().strip()
            priority = str(req.get("priority", "")).lower().strip()

            values = [
                safe_str(req.get("id")),
                _TYPE_LABELS.get(req_type.lower(), req_type) if req_type else "—",
                safe_str(req.get("module")),
                safe_str(req.get("description")),
                _MOSCOW_LABELS.get(priority, priority.capitalize()) if priority else "—",
                safe_str(req.get("source")),
                safe_str(req.get("effort")),
            ]

            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if col_idx in (0, 1, 4, 6):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Цветовое кодирование ячейки приоритета
            priority_cell = row.cells[4]
            if priority in _MOSCOW_CELL_COLORS_HEX:
                self._set_cell_shading(priority_cell, _MOSCOW_CELL_COLORS_HEX[priority])
                text_color = _MOSCOW_TEXT_COLORS_WORD.get(
                    priority, RGBColor(0x33, 0x33, 0x33)
                )
                for p in priority_cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = text_color

        self._format_table_rows(table)

    # ------------------------------------------------------------------
    # Excel-документ
    # ------------------------------------------------------------------

    def generate_excel(
        self,
        requirements: list[dict[str, Any]],
        project_name: str,
        output_path: Path,
    ) -> Path:
        """Генерирует Excel-документ с требованиями.

        Args:
            requirements: Список словарей с полями: id, type, module,
                description, priority, source, effort.
            project_name: Название проекта.
            output_path: Путь для сохранения .xlsx файла.

        Returns:
            Путь к созданному документу.

        Raises:
            ExportError: При ошибке генерации документа.
        """
        try:
            self._ensure_parent_dir(output_path)
            wb = self._create_workbook()

            # Лист «Требования»
            ws_req = wb.active
            ws_req.title = "Требования"
            self._fill_requirements_sheet(ws_req, requirements, project_name)

            # Лист «Сводка»
            ws_summary = wb.create_sheet("Сводка")
            self._fill_summary_sheet(ws_summary, requirements, project_name)

            wb.save(str(output_path))
            logger.info("Документ требований (Excel) сохранён: %s", output_path)
            return output_path

        except Exception as exc:
            logger.exception("Ошибка генерации Excel-документа требований: %s", exc)
            raise ExportError(
                "Ошибка при генерации документа требований (Excel)",
                detail=str(exc),
            ) from exc

    def _fill_requirements_sheet(
        self,
        ws: Any,
        requirements: list[dict[str, Any]],
        project_name: str,
    ) -> None:
        """Заполняет лист «Требования» данными.

        Args:
            ws: Рабочий лист Excel.
            requirements: Список требований.
            project_name: Название проекта.
        """
        # Заголовок документа
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(_EXCEL_HEADERS))
        title_cell = ws.cell(row=1, column=1, value=f"Требования к системе — {project_name}")
        title_cell.font = EXCEL_FONTS["title"]
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        # Дата
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(_EXCEL_HEADERS))
        date_cell = ws.cell(row=2, column=1, value=format_date_russian())
        date_cell.font = Font(name="Calibri", size=10, italic=True, color="FF999999")
        date_cell.alignment = Alignment(horizontal="center")

        # Заголовки таблицы (строка 4)
        header_row = 4
        self._write_excel_header(ws, _EXCEL_HEADERS, row=header_row)

        # Данные
        for row_idx, req in enumerate(requirements, start=header_row + 1):
            req_type = str(req.get("type", "")).upper().strip()
            priority = str(req.get("priority", "")).lower().strip()

            values = [
                safe_str(req.get("id")),
                _TYPE_LABELS.get(req_type.lower(), req_type) if req_type else "—",
                safe_str(req.get("module")),
                safe_str(req.get("description")),
                _MOSCOW_LABELS.get(priority, priority.capitalize()) if priority else "—",
                safe_str(req.get("source")),
                safe_str(req.get("effort")),
            ]

            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = EXCEL_FONTS["normal"]
                cell.border = EXCEL_BORDER

                if col_idx in (1, 2, 5, 7):
                    cell.alignment = EXCEL_ALIGNMENT_CENTER
                else:
                    cell.alignment = EXCEL_ALIGNMENT_LEFT

                # Чередование строк
                if (row_idx - header_row) % 2 == 0:
                    cell.fill = EXCEL_FILLS["alt_row"]

            # Условное форматирование приоритета (столбец 5)
            priority_cell = ws.cell(row=row_idx, column=5)
            if priority in _MOSCOW_FILLS:
                priority_cell.fill = _MOSCOW_FILLS[priority]
                priority_cell.font = Font(
                    name="Calibri",
                    size=11,
                    bold=True,
                    color=_MOSCOW_FONT_COLORS.get(priority, "FF333333"),
                )

        # Автофильтр
        self._add_autofilter(ws, row=header_row)

        # Закрепление областей
        self._freeze_panes(ws, f"A{header_row + 1}")

        # Ширина столбцов
        self._auto_adjust_columns(ws)

        # Принудительные минимальные ширины для ключевых столбцов
        ws.column_dimensions["A"].width = 12  # ID
        ws.column_dimensions["D"].width = 50  # Описание

    def _fill_summary_sheet(
        self,
        ws: Any,
        requirements: list[dict[str, Any]],
        project_name: str,
    ) -> None:
        """Заполняет лист «Сводка» сводными данными по модулям и приоритетам.

        Args:
            ws: Рабочий лист Excel.
            requirements: Список требований.
            project_name: Название проекта.
        """
        # Заголовок
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
        title_cell = ws.cell(row=1, column=1, value=f"Сводка требований — {project_name}")
        title_cell.font = EXCEL_FONTS["title"]
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        # Сбор данных по модулям
        modules: dict[str, dict[str, int]] = {}
        for req in requirements:
            module = safe_str(req.get("module"), "Не указан")
            priority = str(req.get("priority", "")).lower().strip()
            if module not in modules:
                modules[module] = {"must": 0, "should": 0, "could": 0, "wont": 0, "total": 0}
            modules[module]["total"] += 1
            if priority in modules[module]:
                modules[module][priority] += 1

        # Таблица «По модулям»
        header_row = 3
        summary_headers = ["Модуль", "Must", "Should", "Could", "Won't", "Итого"]
        self._write_excel_header(ws, summary_headers, row=header_row)

        for row_idx, (module_name, counts) in enumerate(
            sorted(modules.items(), key=lambda x: x[1]["total"], reverse=True),
            start=header_row + 1,
        ):
            values = [
                module_name,
                counts["must"],
                counts["should"],
                counts["could"],
                counts["wont"],
                counts["total"],
            ]
            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = EXCEL_FONTS["normal"]
                cell.border = EXCEL_BORDER
                if col_idx == 1:
                    cell.alignment = EXCEL_ALIGNMENT_LEFT
                else:
                    cell.alignment = EXCEL_ALIGNMENT_CENTER

                # Чередование строк
                if (row_idx - header_row) % 2 == 0:
                    cell.fill = EXCEL_FILLS["alt_row"]

            # Цветовое кодирование ячеек Must/Should/Could/Won't
            moscow_keys = ["must", "should", "could", "wont"]
            for i, mk in enumerate(moscow_keys):
                c = ws.cell(row=row_idx, column=i + 2)
                if counts[mk] > 0:
                    c.fill = _MOSCOW_FILLS[mk]
                    c.font = Font(
                        name="Calibri",
                        size=11,
                        bold=True,
                        color=_MOSCOW_FONT_COLORS.get(mk, "FF333333"),
                    )

        # Строка итогов
        if modules:
            total_row = header_row + len(modules) + 1
            ws.cell(row=total_row, column=1, value="ИТОГО").font = EXCEL_FONTS["bold"]
            ws.cell(row=total_row, column=1).border = EXCEL_BORDER

            moscow_keys = ["must", "should", "could", "wont"]
            for i, mk in enumerate(moscow_keys):
                total_val = sum(m[mk] for m in modules.values())
                cell = ws.cell(row=total_row, column=i + 2, value=total_val)
                cell.font = EXCEL_FONTS["bold"]
                cell.alignment = EXCEL_ALIGNMENT_CENTER
                cell.border = EXCEL_BORDER

            grand_total = sum(m["total"] for m in modules.values())
            cell_gt = ws.cell(row=total_row, column=6, value=grand_total)
            cell_gt.font = EXCEL_FONTS["bold"]
            cell_gt.alignment = EXCEL_ALIGNMENT_CENTER
            cell_gt.border = EXCEL_BORDER

        # Автофильтр и ширины
        self._add_autofilter(ws, row=header_row)
        self._freeze_panes(ws, f"A{header_row + 1}")
        self._auto_adjust_columns(ws)
        ws.column_dimensions["A"].width = 30  # Модуль
