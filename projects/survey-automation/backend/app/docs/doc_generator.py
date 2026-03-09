"""Базовый генератор документов.

Предоставляет абстрактный базовый класс DocGenerator с общими утилитами
для создания Word- и Excel-документов: стили, колонтитулы, форматирование дат,
метаданные документа.
"""

from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

# Русские названия месяцев в родительном падеже
_MONTH_NAMES_GENITIVE: dict[int, str] = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}

# Общие цветовые константы
COLORS = {
    "critical": RGBColor(0xCC, 0x00, 0x00),
    "high": RGBColor(0xFF, 0x66, 0x00),
    "medium": RGBColor(0xFF, 0xCC, 0x00),
    "low": RGBColor(0x33, 0x99, 0x33),
    "must": RGBColor(0xCC, 0x00, 0x00),
    "should": RGBColor(0xFF, 0x66, 0x00),
    "could": RGBColor(0xFF, 0xCC, 0x00),
    "wont": RGBColor(0x99, 0x99, 0x99),
    "header_bg": RGBColor(0x2F, 0x54, 0x96),
    "header_text": RGBColor(0xFF, 0xFF, 0xFF),
    "alt_row": RGBColor(0xF2, 0xF2, 0xF2),
    "border": RGBColor(0xBF, 0xBF, 0xBF),
}

# Excel-эквиваленты цветов (строки без #)
EXCEL_FILLS = {
    "critical": PatternFill(start_color="FFCC0000", end_color="FFCC0000", fill_type="solid"),
    "high": PatternFill(start_color="FFFF6600", end_color="FFFF6600", fill_type="solid"),
    "medium": PatternFill(start_color="FFFFCC00", end_color="FFFFCC00", fill_type="solid"),
    "low": PatternFill(start_color="FF339933", end_color="FF339933", fill_type="solid"),
    "must": PatternFill(start_color="FFCC0000", end_color="FFCC0000", fill_type="solid"),
    "should": PatternFill(start_color="FFFF6600", end_color="FFFF6600", fill_type="solid"),
    "could": PatternFill(start_color="FFFFCC00", end_color="FFFFCC00", fill_type="solid"),
    "wont": PatternFill(start_color="FF999999", end_color="FF999999", fill_type="solid"),
    "header": PatternFill(start_color="FF2F5496", end_color="FF2F5496", fill_type="solid"),
    "alt_row": PatternFill(start_color="FFF2F2F2", end_color="FFF2F2F2", fill_type="solid"),
    "green": PatternFill(start_color="FF339933", end_color="FF339933", fill_type="solid"),
    "yellow": PatternFill(start_color="FFFFCC00", end_color="FFFFCC00", fill_type="solid"),
    "red": PatternFill(start_color="FFCC0000", end_color="FFCC0000", fill_type="solid"),
}

EXCEL_FONTS = {
    "header": Font(name="Calibri", size=11, bold=True, color="FFFFFFFF"),
    "bold": Font(name="Calibri", size=11, bold=True),
    "normal": Font(name="Calibri", size=11),
    "title": Font(name="Calibri", size=14, bold=True),
}

EXCEL_BORDER = Border(
    left=Side(style="thin", color="FFBFBFBF"),
    right=Side(style="thin", color="FFBFBFBF"),
    top=Side(style="thin", color="FFBFBFBF"),
    bottom=Side(style="thin", color="FFBFBFBF"),
)

EXCEL_ALIGNMENT_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
EXCEL_ALIGNMENT_LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)


def format_date_russian(dt: datetime | None = None) -> str:
    """Форматирует дату в русском стиле: «1 марта 2026 г.»

    Args:
        dt: Дата для форматирования. Если None, используется текущая дата.

    Returns:
        Строка с датой в формате «день месяц год г.».
    """
    if dt is None:
        dt = datetime.now()
    month_name = _MONTH_NAMES_GENITIVE.get(dt.month, "")
    return f"{dt.day} {month_name} {dt.year} г."


def safe_str(value: Any, default: str = "—") -> str:
    """Безопасно преобразует значение в строку.

    Args:
        value: Произвольное значение.
        default: Значение по умолчанию, если value пуст.

    Returns:
        Строковое представление или значение по умолчанию.
    """
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def safe_list(value: Any, default: str = "—") -> str:
    """Преобразует список в строку через запятую.

    Args:
        value: Список или единичное значение.
        default: Значение по умолчанию для пустого списка.

    Returns:
        Строка с перечислением элементов.
    """
    if not value:
        return default
    if isinstance(value, list):
        items = [str(item).strip() for item in value if item]
        return ", ".join(items) if items else default
    return safe_str(value, default)


class DocGenerator(ABC):
    """Абстрактный базовый класс для генераторов документов.

    Предоставляет общие утилиты для создания Word-документов:
    стили, колонтитулы, метаданные, форматирование дат.

    Attributes:
        author: Имя автора документа.
        company: Название организации.
    """

    def __init__(
        self,
        author: str = "Survey Automation",
        company: str = "",
    ) -> None:
        self.author = author
        self.company = company

    # ------------------------------------------------------------------
    # Word-документ: создание и настройка
    # ------------------------------------------------------------------

    def _create_document(self, title: str = "") -> Document:
        """Создаёт новый Word-документ с базовыми настройками.

        Устанавливает формат A4, поля 2 см, метаданные и стили.

        Args:
            title: Заголовок документа для метаданных.

        Returns:
            Настроенный экземпляр Document.
        """
        doc = Document()

        # Метаданные
        core = doc.core_properties
        core.author = self.author
        core.title = title
        core.language = "ru-RU"

        # Формат A4, поля 2 см
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.orientation = WD_ORIENT.PORTRAIT

        self._create_styles(doc)
        return doc

    def _create_styles(self, doc: Document) -> None:
        """Создаёт или обновляет пользовательские стили документа.

        Настраивает стили заголовков (Heading 1-3), основного текста (Normal),
        а также стили таблиц.

        Args:
            doc: Документ, в котором создаются стили.
        """
        styles = doc.styles

        # Normal
        style_normal = styles["Normal"]
        font = style_normal.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = style_normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        # Heading 1
        style_h1 = styles["Heading 1"]
        font_h1 = style_h1.font
        font_h1.name = "Calibri"
        font_h1.size = Pt(18)
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        pf_h1 = style_h1.paragraph_format
        pf_h1.space_before = Pt(24)
        pf_h1.space_after = Pt(12)
        pf_h1.keep_with_next = True

        # Heading 2
        style_h2 = styles["Heading 2"]
        font_h2 = style_h2.font
        font_h2.name = "Calibri"
        font_h2.size = Pt(14)
        font_h2.bold = True
        font_h2.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        pf_h2 = style_h2.paragraph_format
        pf_h2.space_before = Pt(18)
        pf_h2.space_after = Pt(6)
        pf_h2.keep_with_next = True

        # Heading 3
        style_h3 = styles["Heading 3"]
        font_h3 = style_h3.font
        font_h3.name = "Calibri"
        font_h3.size = Pt(12)
        font_h3.bold = True
        font_h3.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        pf_h3 = style_h3.paragraph_format
        pf_h3.space_before = Pt(12)
        pf_h3.space_after = Pt(6)
        pf_h3.keep_with_next = True

    def _add_header(self, doc: Document, text: str) -> None:
        """Добавляет верхний колонтитул с текстом.

        Args:
            doc: Целевой документ.
            text: Текст верхнего колонтитула.
        """
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = "Calibri"

    def _add_footer_with_pages(self, doc: Document) -> None:
        """Добавляет нижний колонтитул с нумерацией страниц.

        Формат: «Стр. X из Y» по центру.

        Args:
            doc: Целевой документ.
        """
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_prefix = paragraph.add_run("Стр. ")
            run_prefix.font.size = Pt(9)
            run_prefix.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_prefix.font.name = "Calibri"

            # Поле PAGE
            fld_page = self._create_field(doc, "PAGE")
            run_page = paragraph.add_run()
            run_page.font.size = Pt(9)
            run_page.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_page._element.append(fld_page)

            run_sep = paragraph.add_run(" из ")
            run_sep.font.size = Pt(9)
            run_sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_sep.font.name = "Calibri"

            # Поле NUMPAGES
            fld_total = self._create_field(doc, "NUMPAGES")
            run_total = paragraph.add_run()
            run_total.font.size = Pt(9)
            run_total.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_total._element.append(fld_total)

    @staticmethod
    def _create_field(doc: Document, field_name: str) -> Any:
        """Создаёт XML-элемент поля Word (PAGE, NUMPAGES и т.п.).

        Args:
            doc: Документ (для доступа к пространству имён).
            field_name: Имя поля Word.

        Returns:
            XML-элемент fldSimple.
        """
        fld = doc.element.makeelement(
            qn("w:fldSimple"),
            {qn("w:instr"): f" {field_name} "},
        )
        return fld

    def _add_title_page(
        self,
        doc: Document,
        title: str,
        project_name: str,
        date_str: str | None = None,
    ) -> None:
        """Добавляет титульную страницу документа.

        Args:
            doc: Целевой документ.
            title: Заголовок документа.
            project_name: Название проекта.
            date_str: Строка с датой. Если None — текущая дата.
        """
        if date_str is None:
            date_str = format_date_russian()

        # Пустые строки для вертикального центрирования
        for _ in range(6):
            doc.add_paragraph("")

        # Заголовок
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title)
        run_title.font.size = Pt(28)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        run_title.font.name = "Calibri"

        doc.add_paragraph("")

        # Название проекта
        p_project = doc.add_paragraph()
        p_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_project = p_project.add_run(project_name)
        run_project.font.size = Pt(16)
        run_project.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_project.font.name = "Calibri"

        doc.add_paragraph("")

        # Дата
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = p_date.add_run(date_str)
        run_date.font.size = Pt(12)
        run_date.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run_date.font.name = "Calibri"

        # Компания, если задана
        if self.company:
            doc.add_paragraph("")
            p_company = doc.add_paragraph()
            p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_company = p_company.add_run(self.company)
            run_company.font.size = Pt(12)
            run_company.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_company.font.name = "Calibri"

        # Разрыв страницы после титула
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Добавляет оглавление (поле TOC, обновляется при открытии Word).

        Args:
            doc: Целевой документ.
        """
        doc.add_heading("Содержание", level=1)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # XML для поля TOC
        fld_char_begin = doc.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char_begin)

        instr_text = doc.element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._element.append(instr_text)

        fld_char_separate = doc.element.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "separate"}
        )
        run._element.append(fld_char_separate)

        # Текст-заглушка
        run_placeholder = paragraph.add_run(
            "Обновите оглавление: щёлкните правой кнопкой → Обновить поле"
        )
        run_placeholder.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run_placeholder.font.size = Pt(10)
        run_placeholder.font.italic = True

        fld_char_end = doc.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run_end = paragraph.add_run()
        run_end._element.append(fld_char_end)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Форматирование таблиц Word
    # ------------------------------------------------------------------

    def _format_table_header(self, table: Any) -> None:
        """Форматирует заголовок таблицы Word (первая строка).

        Применяет синий фон и белый текст к первой строке таблицы.

        Args:
            table: Таблица Word (docx.table.Table).
        """
        if not table.rows:
            return
        for cell in table.rows[0].cells:
            shading = cell._element.makeelement(
                qn("w:shd"),
                {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "2F5496",
                },
            )
            cell._element.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = COLORS["header_text"]
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    def _format_table_rows(self, table: Any, start_row: int = 1) -> None:
        """Применяет чередующуюся заливку строк таблицы Word.

        Args:
            table: Таблица Word.
            start_row: Номер первой строки данных (0 — заголовок).
        """
        for idx, row in enumerate(table.rows[start_row:], start=start_row):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
                if idx % 2 == 0:
                    shading = cell._element.makeelement(
                        qn("w:shd"),
                        {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): "F2F2F2",
                        },
                    )
                    cell._element.get_or_add_tcPr().append(shading)

    def _set_cell_shading(self, cell: Any, color_hex: str) -> None:
        """Устанавливает фоновый цвет ячейки Word.

        Args:
            cell: Ячейка таблицы Word.
            color_hex: Цвет в формате HEX без # (например, «CC0000»).
        """
        shading = cell._element.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): color_hex,
            },
        )
        cell._element.get_or_add_tcPr().append(shading)

    # ------------------------------------------------------------------
    # Excel: общие утилиты
    # ------------------------------------------------------------------

    @staticmethod
    def _create_workbook() -> Workbook:
        """Создаёт новую рабочую книгу Excel.

        Returns:
            Пустая рабочая книга (Workbook).
        """
        wb = Workbook()
        return wb

    @staticmethod
    def _write_excel_header(ws: Any, headers: list[str], row: int = 1) -> None:
        """Записывает заголовочную строку в лист Excel с форматированием.

        Args:
            ws: Рабочий лист (Worksheet).
            headers: Список заголовков столбцов.
            row: Номер строки заголовка (1-based).
        """
        for col_idx, header_text in enumerate(headers, start=1):
            cell = ws.cell(row=row, column=col_idx, value=header_text)
            cell.font = EXCEL_FONTS["header"]
            cell.fill = EXCEL_FILLS["header"]
            cell.alignment = EXCEL_ALIGNMENT_CENTER
            cell.border = EXCEL_BORDER

    @staticmethod
    def _auto_adjust_columns(ws: Any, min_width: float = 10, max_width: float = 50) -> None:
        """Автоматически подгоняет ширину столбцов по содержимому.

        Args:
            ws: Рабочий лист (Worksheet).
            min_width: Минимальная ширина столбца.
            max_width: Максимальная ширина столбца.
        """
        for col_cells in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col_cells[0].column)
            for cell in col_cells:
                if cell.value is not None:
                    cell_length = len(str(cell.value))
                    max_length = max(max_length, cell_length)
            adjusted_width = min(max(max_length + 2, min_width), max_width)
            ws.column_dimensions[col_letter].width = adjusted_width

    @staticmethod
    def _add_autofilter(ws: Any, row: int = 1) -> None:
        """Добавляет автофильтр на указанную строку.

        Args:
            ws: Рабочий лист.
            row: Номер строки с заголовками.
        """
        if ws.max_column and ws.max_row and ws.max_row > row:
            last_col_letter = get_column_letter(ws.max_column)
            ws.auto_filter.ref = f"A{row}:{last_col_letter}{ws.max_row}"

    @staticmethod
    def _freeze_panes(ws: Any, cell: str = "A2") -> None:
        """Закрепляет области в Excel (по умолчанию — первую строку).

        Args:
            ws: Рабочий лист.
            cell: Ячейка, выше и левее которой области закреплены.
        """
        ws.freeze_panes = cell

    @staticmethod
    def _apply_border_to_range(
        ws: Any,
        start_row: int,
        end_row: int,
        start_col: int,
        end_col: int,
    ) -> None:
        """Применяет тонкие границы к диапазону ячеек.

        Args:
            ws: Рабочий лист.
            start_row: Начальная строка.
            end_row: Конечная строка.
            start_col: Начальный столбец.
            end_col: Конечный столбец.
        """
        for row_idx in range(start_row, end_row + 1):
            for col_idx in range(start_col, end_col + 1):
                ws.cell(row=row_idx, column=col_idx).border = EXCEL_BORDER

    # ------------------------------------------------------------------
    # Сохранение
    # ------------------------------------------------------------------

    @staticmethod
    def _ensure_parent_dir(path: Path) -> Path:
        """Создаёт родительскую директорию, если она не существует.

        Args:
            path: Путь к файлу.

        Returns:
            Тот же путь (для удобства цепочек вызовов).
        """
        path.parent.mkdir(parents=True, exist_ok=True)
        return path

    @abstractmethod
    def generate(self, *args: Any, **kwargs: Any) -> Path:
        """Генерирует документ. Реализуется в подклассах.

        Returns:
            Путь к созданному файлу.
        """
        ...
