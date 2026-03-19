"""Generate Skills README as a formatted Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    return table


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_heading("Claude Code Skills — Справочник", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Глобальные скиллы для всех проектов\nВызываются через / в чате Claude Code")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ── Overview table ──
    add_styled_heading(doc, "Обзор всех скиллов", level=1)

    skills_table = [
        ["/full-review", "📋 Мастер", "Запускает ВСЕ проверки и выдаёт единый отчёт с приоритетами"],
        ["/architecture-review", "🏗️ Архитектура", "Аудит структуры проекта, модулей, зависимостей, потока данных"],
        ["/find-bugs", "🐛 Баги", "Поиск runtime-ошибок, логических багов, проблем безопасности"],
        ["/smart-fix", "🔧 Исправление", "Автоматическое исправление найденных багов по приоритету"],
        ["/ui-ux-review", "🎨 UI/UX", "Аудит интерфейса: юзабилити, доступность, пользовательский путь"],
        ["/tdd", "🧪 Тесты", "Генерация unit-тестов по методологии TDD (red → green → refactor)"],
        ["/e2e-test", "🌐 E2E", "End-to-end тесты через Playwright в браузере"],
        ["/lint-check", "✨ Качество", "Линтинг (ruff) + проверка типов (pyright) + автоформатирование"],
        ["/debug", "🔍 Отладка", "Пошаговая отладка: воспроизведение → причина → фикс"],
        ["/health-check", "💚 Диагностика", "Проверка: зависимости, сервисы, конфиг, импорты"],
        ["/1c-erp-development", "🏢 1С", "Правила разработки для 1С:ERP УСО 2.5"],
        ["/detailed-instruction", "📖 Инструкция", "Детальные инструкции по работе с проектами"],
    ]
    add_table(doc, ["Команда", "Категория", "Описание"], skills_table)

    doc.add_page_break()

    # ── Usage scenarios ──
    add_styled_heading(doc, "Сценарии использования", level=1)

    scenarios = [
        ("🚀 Быстрый старт нового проекта",
         "/health-check → /architecture-review → /find-bugs → /smart-fix"),
        ("📋 Полный аудит проекта",
         "/full-review  (одна команда запускает всё)"),
        ("🛠️ Разработка новой фичи",
         "/tdd → написать код → /lint-check"),
        ("🔥 Проблема в проде",
         "/debug <описание> → /find-bugs → /smart-fix"),
        ("🎨 Работа над UI",
         "/ui-ux-review → /e2e-test"),
    ]
    for title_text, commands in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(11)
        p2 = doc.add_paragraph(commands)
        p2.paragraph_format.left_indent = Cm(1)
        for run in p2.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    doc.add_paragraph()

    # ── Detailed descriptions ──
    add_styled_heading(doc, "Детальное описание каждого скилла", level=1)

    details = [
        {
            "cmd": "/full-review",
            "title": "Полный ревью проекта",
            "when": "Перед релизом, при приёмке проекта, для общей оценки состояния.",
            "what": "Последовательно запускает: health-check → architecture-review → find-bugs → ui-ux-review → lint-check. Результат — единый отчёт с оценками по 10-балльной шкале и планом действий на 3 недели.",
        },
        {
            "cmd": "/architecture-review",
            "title": "Ревью архитектуры",
            "when": "В начале работы над проектом, при рефакторинге, при добавлении крупной фичи.",
            "what": "Анализирует структуру проекта и модулей, граф зависимостей (ищет циклические), поток данных от входа до выхода, контракты между модулями, конфигурацию и дефолты, масштабируемость (память, таймауты). Выдаёт оценку X/10.",
        },
        {
            "cmd": "/find-bugs",
            "title": "Поиск багов",
            "when": "Регулярно, перед релизом, после крупных изменений.",
            "what": "Ищет по 5 категориям: (1) Runtime — необработанные исключения, None-доступы, утечки ресурсов. (2) Логические — ошибки условий, off-by-one, несовпадение дефолтов. (3) Интеграция — таймауты API, ошибки авторизации. (4) Безопасность — path traversal, инъекции, секреты в коде. (5) Платформа — пути Windows/Linux, кодировки. Каждый баг получает ID и приоритет.",
        },
        {
            "cmd": "/smart-fix",
            "title": "Автоисправление багов",
            "when": "После /find-bugs, когда нужно быстро исправить найденное.",
            "what": "Берёт результаты /find-bugs и автоматически исправляет: сначала критические (крэши, потеря данных), затем высокий приоритет. Минимальные изменения, не ломающие существующий код. Добавляет обработку ошибок и type hints.",
        },
        {
            "cmd": "/ui-ux-review",
            "title": "UI/UX аудит",
            "when": "После создания UI, перед демонстрацией заказчику.",
            "what": "Проверяет по 10 эвристикам Нильсена: видимость состояния, терминология, возможность отмены, консистентность, предотвращение ошибок, подсказки, доступность (контраст, размеры, Tab-навигация). Выдаёт постраничный разбор + быстрые улучшения.",
        },
        {
            "cmd": "/tdd",
            "title": "Test-Driven Development",
            "when": "При разработке новых модулей, при рефакторинге.",
            "what": "Создаёт полный набор тестов: conftest.py с фикстурами, unit-тесты для каждого модуля, моки для внешних зависимостей (API, ML-модели), интеграционные тесты. Запускает pytest и сообщает результаты.",
        },
        {
            "cmd": "/e2e-test",
            "title": "E2E тестирование",
            "when": "Для автоматизации проверки веб-интерфейса.",
            "what": "Создаёт E2E тесты через Playwright: автозапуск веб-сервера, тест каждой страницы, навигации, сохранения состояния, полный пользовательский путь. Скриншоты при ошибках.",
        },
        {
            "cmd": "/lint-check",
            "title": "Линтинг и типы",
            "when": "После каждого значительного изменения кода.",
            "what": "Запускает ruff (стиль, неиспользуемые импорты, типичные ошибки, автоисправление), ruff format (единообразное форматирование), pyright (статическая проверка типов). Создаёт конфиг ruff если отсутствует.",
        },
        {
            "cmd": "/debug",
            "title": "Систематическая отладка",
            "when": "При конкретной ошибке. Передай описание проблемы как аргумент.",
            "what": "Процесс: (1) воспроизведение ошибки, (2) изоляция — трассировка потока данных до точки сбоя, (3) анализ корневой причины (почему, не что), (4) минимальное исправление, (5) верификация.",
        },
        {
            "cmd": "/health-check",
            "title": "Диагностика системы",
            "when": "При первом запуске, после установки, когда что-то не работает.",
            "what": "Проверяет версию Python и пакеты, доступность внешних сервисов (Ollama, API, БД), системные утилиты (FFmpeg, Node.js), валидность конфигурации, структуру директорий, импорт всех модулей. Выдаёт таблицу: ЗДОРОВ / ЧАСТИЧНО / СЛОМАН.",
        },
    ]

    for d in details:
        add_styled_heading(doc, f"{d['cmd']}  —  {d['title']}", level=2)

        p = doc.add_paragraph()
        run = p.add_run("Когда использовать: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(d["when"])
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run("Что делает: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(d["what"])
        run.font.size = Pt(10)

    doc.add_page_break()

    # ── File structure ──
    add_styled_heading(doc, "Расположение файлов скиллов", level=1)

    p = doc.add_paragraph("Все скиллы хранятся в глобальной папке:")
    p = doc.add_paragraph()
    run = p.add_run("C:\\CLOUDE_PR\\.claude\\skills\\")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph("Эта папка глобальная — скиллы доступны во всех проектах внутри ")
    run = p.add_run("C:\\CLOUDE_PR\\")
    run.font.name = "Consolas"
    run.bold = True

    files = [
        "README.md — этот документ (Markdown-версия)",
        "architecture-review.md",
        "find-bugs.md",
        "smart-fix.md",
        "ui-ux-review.md",
        "tdd.md",
        "e2e-test.md",
        "lint-check.md",
        "debug.md",
        "health-check.md",
        "full-review.md",
        "1c-erp-development.md",
        "detailed-instruction.md",
    ]
    for f in files:
        p = doc.add_paragraph(f, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # Save
    output = Path(r"C:\CLOUDE_PR\.claude\skills\Claude_Code_Skills_Guide.docx")
    doc.save(str(output))
    print(f"Saved: {output}")


if __name__ == "__main__":
    main()
